"""
╔══════════════════════════════════════════════════════════════════╗
║        SCHULBUCHVERWALTUNG  –  Streamlit + Firebase              ║
║        E-Mail/Passwort-Login  |  Firestore Datenbank             ║
║        3 MODI: EINZELN / DOPPEL / FLEXIBEL                       ║
╚══════════════════════════════════════════════════════════════════╝

Datenmodell pro Buch:
  isbn            – Primärschlüssel (String)
  titel           – Buchtitel
  fach            – Schulfach
  klasse          – zugeordnete Klasse(n), z.B. "5/6" oder "5a,6b,7c"
  modus           – "einzeln" / "doppel" / "flexibel"
  
  MODUS 1 - EINZELJAHRGANG:
    umlauf_klassen – dict {klasse: anzahl}
    verfuegbar_next = sum(umlauf_klassen) + max(lager - 5, 0)
  
  MODUS 2 - DOPPELJAHRGANG:
    jahrgang1_klassen – dict {klasse: anzahl} (behalten Bücher)
    jahrgang2_klassen – dict {klasse: anzahl} (geben zurück)
    verfuegbar_next = sum(jahrgang2) + max(lager - 5, 0)
  
  MODUS 3 - FLEXIBLER UMLAUF:
    flex_klassen – dict {klasse: {"umlauf": 12, "zurueck": 5}}
    verfuegbar_next = sum(alle "zurueck") + max(lager - 5, 0)
  
  lager           – Anzahl im Lager
  bedarf_next     – erwartete Schülerzahl nächstes Jahr
  anschaffung     – Datum der Anschaffung (String YYYY-MM-DD)
  bestellbar      – True/False (im Schulbuchkatalog verfügbar)
  notizen         – Freitext
"""

# ── Imports ───────────────────────────────────────────────────────────────────
import streamlit as st
import pandas as pd
import json
import re
import requests
from datetime import datetime, date

import firebase_admin
from firebase_admin import credentials, firestore
import plotly.express as px
from fpdf import FPDF

# DOCX für Bestellscheine
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64

# ═══════════════════════════════════════════════════════════════════════════════
#  KONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

APP_TITLE = "📚 Schulbuchverwaltung"
RESERVE   = 5   # Mindestbestand im Lager

ALLE_KLASSEN = [
    "5a","5b","5c",
    "6a","6b","6c",
    "7a","7b","7c",
    "8a","8b","8c",
    "9a","9b","9c",
    "10a","10b","10c","10g1","10g2",
    "11/1","11/2",
    "12/1","12/2",
]

ALLE_FAECHER = [
    "Mathematik","Deutsch","Englisch","Französisch","Latein","Spanisch",
    "Physik","Chemie","Biologie","Mensch-Natur-Technik","Geographie","Geschichte","Politik",
    "Ethik","Religion","Musik","Kunst","Sport","Informatik","Wirtschaft",
    "Sonstiges",
]

# ═══════════════════════════════════════════════════════════════════════════════
#  FIREBASE
# ═══════════════════════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner="Verbinde mit Firebase …")
def init_firebase():
    """Firebase Admin SDK initialisieren. Credentials aus st.secrets."""
    try:
        if firebase_admin._apps:
            return firestore.client()
        key_dict = dict(st.secrets["firebase"])
        if "private_key" in key_dict:
            key_dict["private_key"] = key_dict["private_key"].replace("\\n", "\n")
        cred = credentials.Certificate(key_dict)
        firebase_admin.initialize_app(cred)
        return firestore.client()
    except Exception as e:
        st.error(f"🔥 Firebase-Fehler: {e}")
        st.stop()


def col_ref(db):
    return db.collection("schulbuecher")


def load_all(db) -> list[dict]:
    """Alle Bücher aus Firestore laden."""
    try:
        docs = col_ref(db).stream()
        result = []
        for d in docs:
            row = d.to_dict()
            row["_id"] = d.id
            
            # Migration: alte Bücher mit doppeljahrgang=True zu modus="doppel"
            if "modus" not in row:
                if row.get("doppeljahrgang", False):
                    row["modus"] = "doppel"
                else:
                    row["modus"] = "einzeln"
            
            # Felder sicherstellen
            if row["modus"] == "doppel":
                if "jahrgang1_klassen" not in row:
                    row["jahrgang1_klassen"] = {}
                if "jahrgang2_klassen" not in row:
                    row["jahrgang2_klassen"] = {}
            elif row["modus"] == "flexibel":
                if "flex_klassen" not in row:
                    row["flex_klassen"] = {}
            else:
                if "umlauf_klassen" not in row:
                    row["umlauf_klassen"] = {}
            
            result.append(row)
        return result
    except Exception as e:
        st.error(f"Ladefehler: {e}")
        return []


def save_book(db, data: dict):
    """Buch anlegen oder überschreiben (ISBN = Dokument-ID)."""
    isbn = str(data.get("isbn","")).strip()
    if not isbn:
        st.error("ISBN darf nicht leer sein!")
        return False
    payload = {k: v for k, v in data.items() if k != "_id"}
    try:
        col_ref(db).document(isbn).set(payload)
        return True
    except Exception as e:
        st.error(f"Speicherfehler: {e}")
        return False


def delete_book(db, isbn: str):
    """Buch unwiderruflich löschen."""
    try:
        col_ref(db).document(isbn).delete()
        return True
    except Exception as e:
        st.error(f"Löschfehler: {e}")
        return False


# ═══════════════════════════════════════════════════════════════════════════════
#  PREIS-ABFRAGE VIA GOOGLE BOOKS API
# ═══════════════════════════════════════════════════════════════════════════════

def hole_preis_von_google_books(isbn: str) -> tuple[float | None, str]:
    """
    Versucht den Preis für eine ISBN von Google Books API zu holen.
    Returns: (preis, quelle) oder (None, "Nicht gefunden")
    """
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q=isbn:{isbn}"
        response = requests.get(url, timeout=5)
        
        if response.status_code != 200:
            return None, "API-Fehler"
        
        data = response.json()
        
        if data.get("totalItems", 0) == 0:
            return None, "Nicht gefunden"
        
        item = data["items"][0]
        sale_info = item.get("saleInfo", {})
        
        # Preis extrahieren
        if sale_info.get("saleability") == "FOR_SALE":
            retail_price = sale_info.get("retailPrice", {})
            if retail_price:
                preis = retail_price.get("amount")
                if preis:
                    return float(preis), "Google Books"
        
        # Fallback: List Price (UVP)
        list_price = sale_info.get("listPrice", {})
        if list_price:
            preis = list_price.get("amount")
            if preis:
                return float(preis), "Google Books (UVP)"
        
        return None, "Kein Preis verfügbar"
        
    except Exception as e:
        return None, f"Fehler: {str(e)}"


def aktualisiere_preise(db, buecher: list[dict]) -> dict:
    """
    Aktualisiert Preise für alle Bücher über Google Books API.
    Returns: {erfolg: int, fehler: int, gesamt: int}
    """
    erfolg = 0
    fehler = 0
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, buch in enumerate(buecher):
        isbn = buch.get("isbn", "")
        if not isbn:
            fehler += 1
            continue
        
        status_text.text(f"Prüfe {i+1}/{len(buecher)}: {buch.get('titel', 'Unbekannt')[:30]}...")
        
        preis, quelle = hole_preis_von_google_books(isbn)
        
        if preis:
            # Preis in Firebase speichern
            try:
                col_ref(db).document(isbn).update({
                    "preis": preis,
                    "preis_quelle": quelle,
                    "preis_aktualisiert": datetime.now().strftime("%Y-%m-%d %H:%M")
                })
                erfolg += 1
            except:
                fehler += 1
        else:
            fehler += 1
        
        progress_bar.progress((i + 1) / len(buecher))
    
    progress_bar.empty()
    status_text.empty()
    
    return {"erfolg": erfolg, "fehler": fehler, "gesamt": len(buecher)}


def importiere_preise_aus_excel(db, buecher: list[dict], uploaded_file) -> dict:
    """
    Importiert Preise aus dem Thüringer Schulbuchkatalog Excel.
    Returns: {erfolg: int, fehler: int, gesamt: int, katalog_jahr: str}
    """
    import io
    
    try:
        # Excel-Datei laden
        file_bytes = uploaded_file.read()
        df_katalog = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, skiprows=2)
        
        # Prüfe ob erforderliche Spalten vorhanden sind
        if 'ISBN' not in df_katalog.columns or 'Preis' not in df_katalog.columns:
            return {"erfolg": 0, "fehler": len(buecher), "gesamt": len(buecher), 
                   "katalog_jahr": "Unbekannt", "error": "ISBN oder Preis-Spalte nicht gefunden"}
        
        # Extrahiere Schuljahr aus dem Katalog (z.B. "2025/26")
        katalog_jahr = "Unbekannt"
        for col in df_katalog.columns:
            if '/' in str(col) and len(str(col)) == 7:  # z.B. "2025/26"
                katalog_jahr = str(col)
                break
        
        # Erstelle ISBN -> Preis Mapping
        isbn_preis_map = {}
        for _, row in df_katalog.iterrows():
            isbn = str(row.get('ISBN', '')).strip()
            preis_str = str(row.get('Preis', '')).strip()
            
            # ISBN normalisieren (entferne Leerzeichen, Bindestriche optional)
            isbn_clean = isbn.replace('-', '').replace(' ', '')
            
            # Preis parsen (z.B. "30.50" oder "30,50")
            try:
                preis = float(preis_str.replace(',', '.'))
                if preis > 0:
                    isbn_preis_map[isbn_clean] = preis
                    # Auch mit Bindestrichen speichern
                    isbn_preis_map[isbn] = preis
            except:
                pass
        
        # Aktualisiere Preise
        erfolg = 0
        fehler = 0
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, buch in enumerate(buecher):
            isbn = buch.get("isbn", "").strip()
            if not isbn:
                fehler += 1
                continue
            
            status_text.text(f"Aktualisiere {i+1}/{len(buecher)}: {buch.get('titel', 'Unbekannt')[:30]}...")
            
            # Versuche ISBN mit und ohne Bindestriche
            isbn_clean = isbn.replace('-', '').replace(' ', '')
            
            preis = isbn_preis_map.get(isbn) or isbn_preis_map.get(isbn_clean)
            
            if preis:
                try:
                    col_ref(db).document(isbn).update({
                        "preis": preis,
                        "preis_quelle": f"Schulbuchkatalog {katalog_jahr}",
                        "preis_aktualisiert": datetime.now().strftime("%Y-%m-%d %H:%M")
                    })
                    erfolg += 1
                except Exception as e:
                    fehler += 1
            else:
                fehler += 1
            
            progress_bar.progress((i + 1) / len(buecher))
        
        progress_bar.empty()
        status_text.empty()
        
        return {
            "erfolg": erfolg, 
            "fehler": fehler, 
            "gesamt": len(buecher),
            "katalog_jahr": katalog_jahr
        }
        
    except Exception as e:
        return {
            "erfolg": 0, 
            "fehler": len(buecher), 
            "gesamt": len(buecher),
            "katalog_jahr": "Unbekannt",
            "error": str(e)
        }


# ═══════════════════════════════════════════════════════════════════════════════
#  KATALOG-LOOKUP FUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

def lade_katalog_cache():
    """Lädt den Schulbuchkatalog und cached ihn."""
    if 'katalog_cache' not in st.session_state:
        try:
            # Versuche verschiedene Pfade
            katalog_pfade = [
                "Katalog.xlsm",  # Im gleichen Verzeichnis (GitHub)
                "./Katalog.xlsm",
                "/mnt/user-data/uploads/Katalog.xlsm",
            ]
            
            df_katalog = None
            for pfad in katalog_pfade:
                try:
                    df_katalog = pd.read_excel(pfad, skiprows=2, engine='openpyxl')
                    break
                except:
                    continue
            
            if df_katalog is None:
                st.session_state['katalog_cache'] = {}
                return {}
            
            # Cache als Dictionary: ISBN -> Buchdaten
            katalog_dict = {}
            for _, row in df_katalog.iterrows():
                isbn = str(row.get('ISBN', '')).strip()
                if isbn and isbn != 'nan':
                    # Normalisiere ISBN (mit und ohne Bindestriche)
                    isbn_clean = isbn.replace('-', '').replace(' ', '')
                    
                    buchdaten = {
                        'titel': str(row.get('Titel', '')),
                        'verlag': str(row.get('Verlag', '')),
                        'preis': str(row.get('Preis', '')),
                        'fach': str(row.get('Fach', ''))
                    }
                    
                    katalog_dict[isbn] = buchdaten
                    katalog_dict[isbn_clean] = buchdaten  # Auch ohne Bindestriche
            
            st.session_state['katalog_cache'] = katalog_dict
            return katalog_dict
        except Exception as e:
            st.session_state['katalog_cache'] = {}
            return {}
    
    return st.session_state['katalog_cache']


def hole_buch_aus_katalog(isbn):
    """Holt Buchdaten aus dem Katalog anhand der ISBN."""
    if not isbn:
        return None
        
    katalog = lade_katalog_cache()
    
    # Normalisiere ISBN (entferne Bindestriche)
    isbn_clean = isbn.replace('-', '').replace(' ', '').strip()
    
    # Suche mit und ohne Bindestriche
    for isbn_variant in [isbn, isbn_clean]:
        if isbn_variant in katalog:
            return katalog[isbn_variant]
    
    return None


# ═══════════════════════════════════════════════════════════════════════════════
#  BERECHNUNGEN - 3 MODI
# ═══════════════════════════════════════════════════════════════════════════════

def berechne_felder(buch: dict) -> dict:
    """Berechnete Felder zu einem Buch-Dict hinzufügen."""
    lager  = int(buch.get("lager", 0))
    bedarf = int(buch.get("bedarf_next", 0))
    modus  = buch.get("modus", "einzeln")
    
    if modus == "doppel":
        # DOPPELJAHRGANGS-LOGIK
        jg1 = buch.get("jahrgang1_klassen", {}) or {}
        jg2 = buch.get("jahrgang2_klassen", {}) or {}
        
        jg1_gesamt = sum(int(v) for v in jg1.values())
        jg2_gesamt = sum(int(v) for v in jg2.values())
        umlauf_gesamt = jg1_gesamt + jg2_gesamt
        gesamt = umlauf_gesamt + lager
        
        # NUR Jahrgang 2 kommt zurück
        reserve_verfuegbar = max(lager - RESERVE, 0)
        verfuegbar_next = jg2_gesamt + reserve_verfuegbar
        
        buch["jahrgang1_gesamt"] = jg1_gesamt
        buch["jahrgang2_gesamt"] = jg2_gesamt
        
    elif modus == "flexibel":
        # FLEXIBLER UMLAUF - NEUE LOGIK
        flex = buch.get("flex_klassen", {}) or {}
        
        umlauf_gesamt = 0
        zurueck_gesamt = 0
        
        for kl, data in flex.items():
            if isinstance(data, dict):
                umlauf_gesamt += int(data.get("umlauf", 0))
                zurueck_gesamt += int(data.get("zurueck", 0))
        
        gesamt = umlauf_gesamt + lager
        
        # NUR die zurückkommenden Bücher
        reserve_verfuegbar = max(lager - RESERVE, 0)
        verfuegbar_next = zurueck_gesamt + reserve_verfuegbar
        
        buch["zurueck_gesamt"] = zurueck_gesamt
        
    else:
        # EINZELJAHRGANGS-LOGIK (Standard)
        uk = buch.get("umlauf_klassen", {}) or {}
        umlauf_gesamt = sum(int(v) for v in uk.values())
        gesamt = umlauf_gesamt + lager
        
        # Alle kommen zurück
        reserve_verfuegbar = max(lager - RESERVE, 0)
        verfuegbar_next = umlauf_gesamt + reserve_verfuegbar
    
    differenz = verfuegbar_next - bedarf
    
    buch["umlauf_gesamt"]   = umlauf_gesamt
    buch["gesamt"]          = gesamt
    buch["verfuegbar_next"] = verfuegbar_next
    buch["differenz"]       = differenz
    buch["alarm"]           = differenz < 0
    
    return buch


def buecher_zu_df(buecher: list[dict]) -> pd.DataFrame:
    """Liste von Buch-Dicts → übersichtlicher DataFrame."""
    rows = []
    for b in buecher:
        b = berechne_felder(b)
        
        modus = b.get("modus", "einzeln")
        
        # Umlauf-Spalte formatieren
        if modus == "doppel":
            jg1 = b.get("jahrgang1_klassen", {}) or {}
            jg2 = b.get("jahrgang2_klassen", {}) or {}
            jg1_str = ", ".join(f"{k}: {v}" for k, v in sorted(jg1.items()))
            jg2_str = ", ".join(f"{k}: {v}" for k, v in sorted(jg2.items()))
            umlauf_str = f"J1: {jg1_str} | J2: {jg2_str}" if jg1_str and jg2_str else jg1_str or jg2_str
        elif modus == "flexibel":
            flex = b.get("flex_klassen", {}) or {}
            parts = []
            for k, data in sorted(flex.items()):
                if isinstance(data, dict):
                    u = data.get("umlauf", 0)
                    z = data.get("zurueck", 0)
                    parts.append(f"{k}: {u}({z}↩)")
            umlauf_str = ", ".join(parts)
        else:
            uk = b.get("umlauf_klassen", {}) or {}
            umlauf_str = ", ".join(f"{k}: {v}" for k, v in sorted(uk.items()))
        
        rows.append({
            "isbn":            b.get("isbn",""),
            "titel":           b.get("titel",""),
            "fach":            b.get("fach",""),
            "klasse":          b.get("klasse",""),
            "modus":           {"einzeln": "Einzeln", "doppel": "Doppel", "flexibel": "Flexibel"}.get(modus, modus),
            "umlauf_klassen":  umlauf_str,
            "umlauf_gesamt":   b.get("umlauf_gesamt", 0),
            "lager":           b.get("lager", 0),
            "gesamt":          b.get("gesamt", 0),
            "bedarf_next":     b.get("bedarf_next", 0),
            "verfuegbar_next": b.get("verfuegbar_next", 0),
            "differenz":       b.get("differenz", 0),
            "alarm":           b.get("alarm", False),
            "preis":           b.get("preis", None),
            "preis_quelle":    b.get("preis_quelle", ""),
            "anschaffung":     b.get("anschaffung",""),
            "bestellbar":      b.get("bestellbar", False),
            "notizen":         b.get("notizen",""),
        })
    return pd.DataFrame(rows)


# ═══════════════════════════════════════════════════════════════════════════════
#  AUTHENTIFIZIERUNG
# ═══════════════════════════════════════════════════════════════════════════════

def check_login(email: str, password: str) -> bool:
    try:
        users = dict(st.secrets.get("users", {}))
        return users.get(email.strip(), None) == password
    except Exception:
        return False


def render_login_page():
    st.set_page_config(page_title=APP_TITLE, page_icon="📚", layout="centered")
    st.title(APP_TITLE)
    st.markdown("---")
    st.subheader("🔐 Bitte anmelden")

    with st.form("login_form"):
        email    = st.text_input("E-Mail-Adresse", placeholder="lehrer@schule.de")
        password = st.text_input("Passwort", type="password")
        submit   = st.form_submit_button("Einloggen", use_container_width=True)

    if submit:
        if check_login(email, password):
            st.session_state["logged_in"] = True
            st.session_state["user_email"] = email
            st.rerun()
        else:
            st.error("❌ E-Mail oder Passwort falsch.")

    st.markdown("---")
    st.caption("Nutzer werden in der `secrets.toml` unter `[users]` verwaltet.")


# ═══════════════════════════════════════════════════════════════════════════════
#  EXPORT-FUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

def export_txt(df: pd.DataFrame) -> bytes:
    """Tabellarischer TXT-Export."""
    lines = []
    ts    = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    lines.append("=" * 120)
    lines.append(f"  SCHULBUCHVERWALTUNG - Export vom {ts}")
    lines.append("=" * 120)

    header = (
        f"{'ISBN':<18} {'Titel':<28} {'Fach':<14} {'Modus':<10} "
        f"{'Umlf':>5} {'Lag':>5} {'Ges':>5} "
        f"{'Bed.':>5} {'Verf':>5} {'Diff':>5} {'Anschaffung':<12} "
        f"{'Bestell.':<10}"
    )
    lines.append(header)
    lines.append("-" * 120)

    for _, r in df.iterrows():
        diff_str = str(int(r['differenz']))
        alarm    = " ⚠" if r['alarm'] else ""
        line = (
            f"{str(r['isbn']):<18} {str(r['titel'])[:27]:<28} "
            f"{str(r['fach'])[:13]:<14} {str(r.get('modus', 'Einzeln'))[:9]:<10} "
            f"{int(r['umlauf_gesamt']):>5} {int(r['lager']):>5} "
            f"{int(r['gesamt']):>5} {int(r['bedarf_next']):>5} "
            f"{int(r['verfuegbar_next']):>5} "
            f"{diff_str:>5}{alarm:<2} "
            f"{str(r['anschaffung']):<12} "
            f"{'Ja' if r['bestellbar'] else 'Nein':<10}"
        )
        lines.append(line)

    lines.append("=" * 120)
    lines.append(f"  Gesamt: {len(df)} Bücher  |  "
                 f"Alarm (Nachbestellen): {df['alarm'].sum()} Bücher")
    lines.append("=" * 120)
    
    # Schülerzahlen anhängen wenn db verfügbar
    # Wird beim Aufruf übergeben
    return "\n".join(lines).encode("utf-8")


def export_schuelerzahlen_txt(db, schuljahr=None) -> str:
    """Exportiert Schülerzahlen als Text."""
    lines = []
    lines.append("")
    lines.append("=" * 120)
    lines.append("  SCHÜLERZAHLEN")
    lines.append("=" * 120)
    
    if not schuljahr:
        # Aktuelles Schuljahr
        from datetime import datetime
        aktuelles_jahr = datetime.now().year
        aktueller_monat = datetime.now().month
        if aktueller_monat >= 8:
            schuljahr_start = aktuelles_jahr
        else:
            schuljahr_start = aktuelles_jahr - 1
        schuljahr_end = schuljahr_start + 1
        schuljahr = f"{schuljahr_start}/{schuljahr_end}"
    
    try:
        schueler_ref = db.collection('schuelerzahlen').document(schuljahr)
        schueler_doc = schueler_ref.get()
        if schueler_doc.exists:
            schueler_data = schueler_doc.to_dict()
            lines.append(f"Schuljahr: {schuljahr}")
            lines.append("-" * 120)
            
            jahrgaenge = ["5", "6", "7", "8", "9", "10", "10g", "11", "12"]
            for jg in jahrgaenge:
                gesamt = schueler_data.get(f"{jg}_gesamt", 0)
                sn = schueler_data.get(f"{jg}_spanisch", 0)
                rel = schueler_data.get(f"{jg}_religion", 0)
                daz = schueler_data.get(f"{jg}_daz", 0)
                
                if gesamt > 0 or sn > 0 or rel > 0 or daz > 0:
                    lines.append(f"Klasse {jg:>3}: {gesamt:>3} Schüler (Sn: {sn:>2}, Rel: {rel:>2}, DaZ: {daz:>2})")
            
            lines.append("=" * 120)
        else:
            lines.append(f"Keine Schülerzahlen für {schuljahr} gefunden.")
            lines.append("=" * 120)
    except Exception as e:
        lines.append(f"Fehler beim Laden: {e}")
        lines.append("=" * 120)
    
    return "\n".join(lines)


def export_pdf(df: pd.DataFrame) -> bytes:
    """PDF-Export mit fpdf2."""
    
    def clean_text(text):
        """Entferne alle Unicode-Zeichen die PDF nicht mag."""
        text = str(text)
        # Ersetze Gedankenstriche und andere problematische Zeichen
        replacements = {
            '\u2013': '-',  # en dash
            '\u2014': '-',  # em dash
            '\u2018': "'",  # left single quote
            '\u2019': "'",  # right single quote
            '\u201c': '"',  # left double quote
            '\u201d': '"',  # right double quote
            'ä': 'ae', 'ö': 'oe', 'ü': 'ue',
            'Ä': 'Ae', 'Ö': 'Oe', 'Ü': 'Ue',
            'ß': 'ss',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Entferne alle restlichen Non-ASCII Zeichen
        return text.encode('ascii', 'ignore').decode('ascii')
    
    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Schulbuchverwaltung - Bestandsliste", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(0, 6, f"Stand: {datetime.now().strftime('%d.%m.%Y %H:%M')}", 
             align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    cols   = ["isbn","titel","fach","modus","umlauf_gesamt","lager",
              "gesamt","bedarf_next","verfuegbar_next","differenz","anschaffung","bestellbar"]
    labels = ["ISBN","Titel","Fach","Modus","Umlauf","Lager",
              "Ges.","Bedarf","Verf.","Diff","Anschaffung","Best."]
    widths = [25, 50, 22, 15, 12, 10, 10, 12, 12, 10, 22, 10]

    # Header
    pdf.set_fill_color(50, 80, 130)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 7)
    for lbl, w in zip(labels, widths):
        pdf.cell(w, 7, lbl, border=1, align="C", fill=True)
    pdf.ln()

    # Daten
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 7)
    for i, (_, row) in enumerate(df.iterrows()):
        fill = i % 2 == 0
        if fill:
            pdf.set_fill_color(235, 240, 250)
        else:
            pdf.set_fill_color(255, 255, 255)

        if row["alarm"]:
            pdf.set_text_color(180, 0, 0)
        else:
            pdf.set_text_color(0, 0, 0)

        values = [
            clean_text(row["isbn"])[:20],
            clean_text(row["titel"])[:35],
            clean_text(row["fach"])[:14],
            clean_text(row.get("modus", "Einzeln"))[:10],
            str(int(row["umlauf_gesamt"])),
            str(int(row["lager"])),
            str(int(row["gesamt"])),
            str(int(row["bedarf_next"])),
            str(int(row["verfuegbar_next"])),
            str(int(row["differenz"])),
            clean_text(row["anschaffung"])[:12],
            "Ja" if row["bestellbar"] else "Nein",
        ]
        for val, w in zip(values, widths):
            pdf.cell(w, 6, val, border=1, fill=fill)
        pdf.ln()
    
    # Neue Seite für Schülerzahlen würde zu komplex sein
    # Lasse es beim TXT-Export
    
    return bytes(pdf.output())


# ═══════════════════════════════════════════════════════════════════════════════
#  BESTELLSCHEIN-GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_border(cell, **kwargs):
    """Setzt Rahmen für Tabellenzellen."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    edge_el.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


# ═══════════════════════════════════════════════════════════════════════════════
#  ETIKETTEN-GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def generiere_etiketten_docx(titel, klasse, fach, anschaffung, start_nr, anzahl, logo_bytes=None):
    """
    Generiert Etiketten für AVERY Zweckform 3424 (12 Etiketten pro A4).
    Format: 105 x 48 mm pro Etikett
    """
    doc = Document()
    
    # Schmale Seitenränder für A4
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)
    
    inventar_text = ("Dieses Schulbuch wurde in das Inventar aufgenommen und soll mehrere Jahre verwendet "
                     "werden können. Darum dürfen keine Einträge, Unterstreichungen oder Markierungen vorgenommen "
                     "werden. Zum Schutz des Buches soll es in einem Umschlag eingeschlagen werden.")
    
    # Erstelle Etiketten
    etiketten_pro_seite = 12
    anzahl_seiten = (anzahl + etiketten_pro_seite - 1) // etiketten_pro_seite
    
    etikett_nr = start_nr
    
    for seite in range(anzahl_seiten):
        # Tabelle mit 4 Zeilen x 3 Spalten = 12 Etiketten
        table = doc.add_table(rows=4, cols=3)
        table.autofit = False
        
        # Spaltenbreiten: 105mm = 4.13 inches
        for col in table.columns:
            col.width = Inches(2.72)  # 69mm (passt besser auf A4)
        
        for row_idx in range(4):
            table.rows[row_idx].height = Inches(1.89)  # 48mm
            
            for col_idx in range(3):
                if etikett_nr > start_nr + anzahl - 1:
                    break
                
                cell = table.rows[row_idx].cells[col_idx]
                
                # Zelle als 2-Spalten Tabelle (links: Text, rechts: Logo)
                inner_table = cell.add_table(rows=1, cols=2)
                inner_table.autofit = False
                inner_table.columns[0].width = Inches(1.8)  # Links: Text
                inner_table.columns[1].width = Inches(0.9)  # Rechts: Logo
                
                left_cell = inner_table.rows[0].cells[0]
                right_cell = inner_table.rows[0].cells[1]
                
                # LINKE SEITE: Buch-Infos
                left_para = left_cell.paragraphs[0]
                left_para.paragraph_format.space_before = Pt(2)
                left_para.paragraph_format.space_after = Pt(0)
                
                # Titel
                run = left_para.add_run(f"Titel: {titel}\n")
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                # Klasse
                run = left_para.add_run(f"Klasse: {klasse}\n")
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                # Fach
                run = left_para.add_run(f"Fach: {fach}\n")
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                # Anschaffung
                run = left_para.add_run(f"Anschaffung: {anschaffung}\n\n")
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                # Nummer
                run = left_para.add_run(f"# {etikett_nr}")
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                
                # RECHTE SEITE: Logo + Text
                right_para = right_cell.paragraphs[0]
                right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                right_para.paragraph_format.space_before = Pt(2)
                
                # Logo einfügen wenn vorhanden
                if logo_bytes:
                    try:
                        from io import BytesIO
                        right_para.add_run().add_picture(BytesIO(logo_bytes), width=Inches(0.6))
                        right_para.add_run("\n")
                    except:
                        pass
                
                # Schulname
                run = right_para.add_run("Lobdeburgschule Jena\n\n")
                run.font.name = 'Arial'
                run.font.size = Pt(7)
                
                # Inventar-Text
                run = right_para.add_run(inventar_text)
                run.font.name = 'Arial'
                run.font.size = Pt(6)
                
                etikett_nr += 1
        
        # Neue Seite nach jeder Etiketten-Seite (außer letzte)
        if seite < anzahl_seiten - 1:
            doc.add_page_break()
    
    # Als bytes zurückgeben
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  FACH-MAPPING UND SORTIERUNG
# ═══════════════════════════════════════════════════════════════════════════════

# Fach-Mapping: Voller Name -> Abkürzung
FACH_ABKUERZUNGEN = {
    "Deutsch": "D",
    "Mathematik": "Ma",
    "Englisch": "Eng",
    "Mensch-Natur-Technik": "MNT",
    "Geographie": "Geo",
    "Geschichte": "Ge",
    "Biologie": "Bio",
    "Chemie": "Ch",
    "Physik": "Ph",
    "Astronomie": "A",
    "Spanisch": "Sn",
    "Ethik": "Eth",
    "Religion": "Rel",
    "Wirtschaft/Recht": "WR",
    "Deutsch als Zweitsprache": "DaZ",
}

# Feste Reihenfolge der Fächer
FACH_REIHENFOLGE = ["D", "Ma", "Eng", "MNT", "Geo", "Ge", "Bio", "Ch", "Ph", "A", "Sn", "Eth", "Rel", "WR", "DaZ"]

def fach_zu_abkuerzung(fach_name):
    """Wandelt Fach-Namen in Abkürzung um."""
    return FACH_ABKUERZUNGEN.get(fach_name, fach_name)

def sortiere_buecher_nach_fach(buecher_liste):
    """Sortiert Bücher nach fester Fach-Reihenfolge."""
    def fach_sortkey(buch):
        fach = buch.get("fach", "")
        abkuerzung = fach_zu_abkuerzung(fach)
        try:
            return FACH_REIHENFOLGE.index(abkuerzung)
        except ValueError:
            return 999  # Unbekannte Fächer ans Ende
    
    return sorted(buecher_liste, key=fach_sortkey)


def sortiere_buecher_fach_jahrgang(buecher_liste):
    """Sortiert Bücher nach Fach, dann nach Jahrgang für Dropdowns."""
    # Jahrgangs-Reihenfolge
    jahrgang_order = ["5", "6", "7", "8", "9", "10", "10g", "11", "12"]
    
    def sort_key(buch):
        fach = buch.get("fach", "")
        klasse = str(buch.get("klasse", "")).strip()
        
        # Fach-Index
        abkuerzung = fach_zu_abkuerzung(fach)
        try:
            fach_idx = FACH_REIHENFOLGE.index(abkuerzung)
        except ValueError:
            fach_idx = 999
        
        # Jahrgang-Index (extrahiere erste Zahl/Ziffer)
        jahrgang_idx = 999
        for jg in jahrgang_order:
            if jg in klasse:
                jahrgang_idx = jahrgang_order.index(jg)
                break
        
        return (fach_idx, jahrgang_idx)
    
    return sorted(buecher_liste, key=sort_key)


def generiere_bestellschein_bytes(klassenstufe, schuljahr, lehrbuecher, 
                                    arbeitshefte, arbeitsheft_leerzeilen,
                                    weiterfuehrung, weiterfuehrung_leerzeilen,
                                    logo_bytes=None):
    """Generiert Bestellschein als DOCX bytes."""
    doc = Document()
    
    # Seitenränder
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # HEADER
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(5.0)
    header_table.columns[1].width = Inches(1.5)
    
    left_cell = header_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    
    for text, size in [("Lobdeburgschule\n", 10), ("Unter der Lobdeburg 04\n", 8), 
                       ("07747 Jena\n", 8), ("Telefon: 03641 / 33 11 48", 8)]:
        run = left_para.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Cambria'  # Einheitlich mit Rest des Dokuments
    
    right_cell = header_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if logo_bytes:
        try:
            logo_stream = io.BytesIO(logo_bytes)
            right_para.add_run().add_picture(logo_stream, width=Inches(0.8))
        except:
            pass
    
    for cell in header_table.rows[0].cells:
        set_cell_border(cell, top={"sz": 0, "val": "none"}, bottom={"sz": 0, "val": "none"},
                       left={"sz": 0, "val": "none"}, right={"sz": 0, "val": "none"})
    
    # TITEL
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Schulbuchbestellschein")
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Schuljahr
    info_para = doc.add_paragraph()
    info_para.add_run(f"Schuljahr {schuljahr}\n").font.size = Pt(8)
    run = info_para.add_run(f"Klassenstufe {klassenstufe}")
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Schüler-Daten
    schueler_table = doc.add_table(rows=2, cols=3)
    schueler_table.autofit = False
    for i, width in enumerate([2.2, 2.2, 2.0]):
        schueler_table.columns[i].width = Inches(width)
    
    schueler_table.rows[0].cells[0].text = "." * 40
    schueler_table.rows[0].cells[1].text = "." * 40
    schueler_table.rows[0].cells[2].text = "." * 40
    
    for cell in schueler_table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    
    for i, label in enumerate(["Name", "Vorname", "künftige Klasse"]):
        cell = schueler_table.rows[1].cells[i]
        cell.text = label
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.size = Pt(8)
    
    for row in schueler_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 0, "val": "none"}, bottom={"sz": 0, "val": "none"},
                           left={"sz": 0, "val": "none"}, right={"sz": 0, "val": "none"})
    
    # Hinweis
    hinweis = doc.add_paragraph("Kopieren Sie bitte den Schulbuchbestellschein. Sie benötigen ihn beim Einkauf.")
    hinweis.runs[0].font.size = Pt(8)
    
    # LEHRBÜCHER
    if lehrbuecher:
        # SORTIERE nach Fach-Reihenfolge!
        lehrbuecher_sortiert = sortiere_buecher_nach_fach(lehrbuecher)
        
        lb_title = doc.add_paragraph("Bei Ausleihe ausfüllen")
        lb_title.runs[0].font.bold = True
        lb_title.paragraph_format.space_after = Pt(2)  # 0,2cm = 2pt
        
        hinweis2 = doc.add_paragraph("Notieren Sie bitte die Registriernummer des Lehrbuches und bewerten Sie den Zustand: neu ++ gut + mittel ○ schlecht −")
        hinweis2.runs[0].font.size = Pt(8)
        hinweis2.runs[0].font.italic = True
        hinweis2.paragraph_format.space_before = Pt(0)  # Kein Abstand davor!
        hinweis2.paragraph_format.space_after = Pt(2)   # 0,2cm danach
        
        table = doc.add_table(rows=1 + len(lehrbuecher_sortiert), cols=10)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Spaltenbreiten - frei/Kauf/vorh MINIMAL!
        table.columns[0].width = Inches(0.4)   # Fach
        table.columns[1].width = Inches(2.3)   # Lehrbuch
        table.columns[2].width = Inches(0.9)   # Verlag
        table.columns[3].width = Inches(1.1)   # ISBN
        table.columns[4].width = Inches(0.45)  # Preis
        table.columns[5].width = Inches(0.2)   # frei
        table.columns[6].width = Inches(0.2)   # Kauf
        table.columns[7].width = Inches(0.2)   # vorh
        table.columns[8].width = Inches(0.75)  # Reg-Nr.
        table.columns[9].width = Inches(0.55)  # Zustand
        
        headers = ["Fach", "Lehrbuch", "Verlag", "ISBN", "Preis", "frei", "Kauf", "vorh", "Reg-Nr.", "Zustand"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = 1
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        for idx, buch in enumerate(lehrbuecher_sortiert):
            row = table.rows[idx + 1]
            
            # Zeilenhöhe auf 0,4cm (11,34 pt)
            row.height = Pt(11.34)
            
            # ABKÜRZUNG verwenden!
            fach_abk = fach_zu_abkuerzung(buch.get("fach", ""))
            row.cells[0].text = fach_abk
            row.cells[1].text = buch.get("titel", "")
            row.cells[2].text = buch.get("verlag", "")
            row.cells[3].text = buch.get("isbn", "")
            row.cells[4].text = buch.get("preis", "")
            
            if buch.get("klassensatz", False):
                row.cells[5].text = ""
                row.cells[6].text = ""
                row.cells[7].text = ""
                for i in [5, 6, 7]:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D9D9D9')
                    row.cells[i]._element.get_or_add_tcPr().append(shading)
                row.cells[8].text = "Klassensatz"
                row.cells[8].paragraphs[0].runs[0].font.size = Pt(7)
                row.cells[8].paragraphs[0].runs[0].font.italic = True
            else:
                row.cells[5].text = ""  # frei
                row.cells[6].text = ""  # Kauf
                row.cells[7].text = ""  # vorh
            
            for cell in row.cells[:9]:
                cell.vertical_alignment = 1
                # Zeilenabstand auf 1 setzen
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
    
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(6)
    
    # ARBEITSHEFTE
    total_ah = len(arbeitshefte) + arbeitsheft_leerzeilen
    if total_ah > 0:
        # SORTIERE nach Fach-Reihenfolge!
        arbeitshefte_sortiert = sortiere_buecher_nach_fach(arbeitshefte)
        
        doc.add_paragraph("Arbeitshefte/Lektüren").runs[0].font.bold = True
        ah_table = doc.add_table(rows=1 + total_ah, cols=5)
        ah_table.style = 'Table Grid'
        
        for i, header in enumerate(["Fach", "Titel", "Verlag", "ISBN", "Preis"]):
            cell = ah_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            cell._element.get_or_add_tcPr().append(shading)
        
        for idx, ah in enumerate(arbeitshefte_sortiert):
            row = ah_table.rows[idx + 1]
            row.height = Pt(11.34)  # 0,4cm
            
            # ABKÜRZUNG verwenden!
            fach_abk = fach_zu_abkuerzung(ah.get("fach", ""))
            row.cells[0].text = fach_abk
            row.cells[1].text = ah.get("titel", "")
            row.cells[2].text = ah.get("verlag", "")
            row.cells[3].text = ah.get("isbn", "")
            row.cells[4].text = ah.get("preis", "")
            for cell in row.cells:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.vertical_alignment = 1
    
    # WEITERFÜHRUNG
    total_wf = len(weiterfuehrung) + weiterfuehrung_leerzeilen
    if total_wf > 0:
        # SORTIERE nach Fach-Reihenfolge!
        weiterfuehrung_sortiert = sortiere_buecher_nach_fach(weiterfuehrung)
        
        aus_klasse = weiterfuehrung_sortiert[0].get("aus_klasse", "5") if weiterfuehrung_sortiert else "5"
        wf_title = doc.add_paragraph(f"Weiterführung aus Klasse {aus_klasse}")
        wf_title.runs[0].font.bold = True
        wf_title.paragraph_format.space_after = Pt(0)  # Kein Abstand danach
        
        wf_hint = doc.add_paragraph("Bitte Registriernummer und Zustand der Lehrbücher vom Vorjahr übernehmen.")
        wf_hint.runs[0].font.size = Pt(8)
        wf_hint.paragraph_format.space_before = Pt(0)  # Kein Abstand davor
        
        wf_table = doc.add_table(rows=1 + total_wf, cols=10)
        wf_table.style = 'Table Grid'
        wf_table.autofit = False
        
        # Spaltenbreiten GLEICH wie Lehrbücher!
        wf_table.columns[0].width = Inches(0.4)   # Fach
        wf_table.columns[1].width = Inches(2.3)   # Lehrbuch
        wf_table.columns[2].width = Inches(0.9)   # Verlag
        wf_table.columns[3].width = Inches(1.1)   # ISBN
        wf_table.columns[4].width = Inches(0.2)   # frei
        wf_table.columns[5].width = Inches(0.2)   # Kauf
        wf_table.columns[6].width = Inches(0.2)   # vorh
        wf_table.columns[7].width = Inches(0.75)  # Reg-Nr.
        wf_table.columns[8].width = Inches(0.55)  # Zustand
        
        headers = ["Fach", "Lehrbuch", "Verlag", "ISBN", "frei", "Kauf", "vorh", "Reg-Nr.", "Zustand"]
        for i, header in enumerate(headers):
            cell = wf_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = 1
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')  # Nur Header grau
            cell._element.get_or_add_tcPr().append(shading)
        
        for idx, wf in enumerate(weiterfuehrung_sortiert):
            row = wf_table.rows[idx + 1]
            row.height = Pt(11.34)  # 0,4cm
            
            # ABKÜRZUNG verwenden!
            fach_abk = fach_zu_abkuerzung(wf.get("fach", ""))
            row.cells[0].text = fach_abk
            row.cells[1].text = wf.get("titel", "")
            row.cells[2].text = wf.get("verlag", "")
            row.cells[3].text = wf.get("isbn", "")
            # frei, Kauf, vorh bleiben leer
            for cell in row.cells[:9]:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.vertical_alignment = 1  # Vertikal zentrieren
    
    # FOOTER
    doc.add_paragraph()  # Leerzeile
    footer = doc.add_paragraph("Bitte erst nach dem Erhalt der Freiexemplare unterschreiben.\n\n")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.italic = True
    
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.columns[0].width = Inches(3.2)
    footer_table.columns[1].width = Inches(3.2)
    
    footer_table.rows[0].cells[0].text = "." * 80 + "\nUnterschrift des Schülers bei Ausleihe"
    
    # Klasse 11/12: Hinweis zu Volljährigkeit
    if klassenstufe in ["11", "12"]:
        footer_table.rows[0].cells[1].text = "." * 80 + "\nUnterschrift der Eltern bei Ausleihe\n(nur bei Minderjährigen)"
    else:
        footer_table.rows[0].cells[1].text = "." * 80 + "\nUnterschrift der Eltern bei Ausleihe"
    
    for cell in footer_table.rows[0].cells:
        set_cell_border(cell, top={"sz": 0, "val": "none"}, bottom={"sz": 0, "val": "none"},
                       left={"sz": 0, "val": "none"}, right={"sz": 0, "val": "none"})
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
    
    # Als bytes zurückgeben
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  FORMULAR MIT 3 MODI: EINZELN / DOPPEL / FLEXIBEL
# ═══════════════════════════════════════════════════════════════════════════════

def hole_schuelerzahl_fuer_buch(db, fach, klasse, schuljahr=None):
    """Holt die passende Schülerzahl für ein Buch aus Firebase."""
    if not schuljahr:
        # Aktuelles Schuljahr berechnen
        from datetime import datetime
        aktuelles_jahr = datetime.now().year
        aktueller_monat = datetime.now().month
        if aktueller_monat >= 8:
            schuljahr_start = aktuelles_jahr
        else:
            schuljahr_start = aktuelles_jahr - 1
        schuljahr_end = schuljahr_start + 1
        schuljahr = f"{schuljahr_start}/{schuljahr_end}"
    
    try:
        schueler_ref = db.collection('schuelerzahlen').document(schuljahr)
        schueler_doc = schueler_ref.get()
        if not schueler_doc.exists:
            return 0
        
        schueler_data = schueler_doc.to_dict()
        
        # Extrahiere Jahrgang aus klasse (z.B. "5" aus "5a" oder "10g" aus "10g.1")
        jahrgang = klasse.strip()
        # Versuche gängige Formate
        if jahrgang and jahrgang[0].isdigit():
            # Extrahiere Zahl + optional "g"
            import re
            match = re.match(r'(\d+g?)', jahrgang)
            if match:
                jahrgang = match.group(1)
        
        # Prüfe ob Spezialfach
        if "Spanisch" in fach or "Sn" in fach:
            return schueler_data.get(f"{jahrgang}_spanisch", 0)
        elif "Religion" in fach or "Rel" in fach:
            return schueler_data.get(f"{jahrgang}_religion", 0)
        elif "DaZ" in fach or "Deutsch als Zweitsprache" in fach:
            return schueler_data.get(f"{jahrgang}_daz", 0)
        else:
            # Standard-Fach: Gesamt-Zahl
            return schueler_data.get(f"{jahrgang}_gesamt", 0)
    except Exception as e:
        return 0


def buch_formular(db, existing: dict | None = None):
    """
    Formular zum Anlegen oder Bearbeiten eines Buches.
    3 MODI: Einzeljahrgang / Doppeljahrgang / Flexibler Umlauf
    """
    is_new = existing is None
    if is_new:
        existing = {
            "isbn":"", "titel":"", "fach": ALLE_FAECHER[0],
            "klasse":"", "modus": "einzeln",
            "umlauf_klassen":{}, "jahrgang1_klassen":{}, "jahrgang2_klassen":{}, "flex_klassen":{},
            "lager":0, "bedarf_next":0,
            "anschaffung": str(date.today()),
            "bestellbar": True, "klassensatz": False, "notizen":"",
        }
    
    # Eindeutige Form-ID basierend auf Kontext (neu vs. edit vs. detail)
    import hashlib
    form_context = f"{'new' if is_new else existing.get('isbn', 'edit')}"
    form_id = hashlib.md5(form_context.encode()).hexdigest()[:8]

    with st.form(key=f"buch_form_{form_id}"):
        st.markdown("### " + ("➕ Neues Buch anlegen" if is_new else f"✏️ Bearbeiten: {existing.get('titel','')}"))

        c1, c2 = st.columns(2)
        with c1:
            isbn_input = st.text_input("ISBN *", value=existing["isbn"],
                                   help="ISBN ist der eindeutige Schlüssel - Daten werden automatisch aus Katalog geladen" + 
                                   ("" if is_new else " ⚠️ Achtung: Änderung erstellt neues Buch!"))
            
            # Auto-Lookup aus Katalog wenn ISBN eingegeben wird
            if isbn_input and isbn_input != existing.get("isbn_last_lookup", ""):
                katalog_data = hole_buch_aus_katalog(isbn_input)
                if katalog_data:
                    # Nur überschreiben wenn Felder noch leer sind
                    if not existing.get("titel"):
                        existing["titel"] = katalog_data.get("titel", "")
                    if not existing.get("verlag"):
                        existing["verlag"] = katalog_data.get("verlag", "")
                    if not existing.get("fach") or existing.get("fach") == ALLE_FAECHER[0]:
                        existing["fach"] = katalog_data.get("fach", "")
                    if not existing.get("preis"):
                        try:
                            preis_str = katalog_data.get("preis", "").replace("€", "").replace(",", ".").strip()
                            if preis_str:
                                existing["preis"] = float(preis_str)
                        except:
                            pass
                    existing["isbn_last_lookup"] = isbn_input
                    if is_new:
                        st.success(f"✅ Daten aus Katalog geladen: {katalog_data.get('titel', '')[:50]}")
            
            isbn = isbn_input
            titel  = st.text_input("Titel *", value=existing.get("titel", ""),
                                   help="Wird automatisch aus Katalog geladen")
            fach_idx = ALLE_FAECHER.index(existing.get("fach", "")) if existing.get("fach", "") in ALLE_FAECHER else 0
            fach   = st.selectbox("Fach", ALLE_FAECHER, index=fach_idx)
            verlag = st.text_input("Verlag", value=existing.get("verlag", ""),
                                   help="Wird automatisch aus Katalog geladen")
            klasse = st.text_input("Jahrgang / Klassenstufe",
                                   value=existing.get("klasse", ""),
                                   help='Nur Jahrgang eingeben, z.B. "5", "7", "10g", "11" - OHNE a,b,c oder .1/.2')
        with c2:
            lager   = st.number_input("Lager (Exemplare)", min_value=0,
                                      value=int(existing.get("lager", 0)))
            
            # Auto-Bedarf aus Schülerzahlen wenn neues Buch
            if is_new and fach and klasse:
                auto_bedarf = hole_schuelerzahl_fuer_buch(db, fach, klasse)
            else:
                auto_bedarf = int(existing.get("bedarf_next", 0))
            
            bedarf  = st.number_input("Bedarf nächstes Jahr (neue Schüler)", min_value=0,
                                      value=auto_bedarf,
                                      help="Anzahl Schüler die nächstes Jahr das Buch NEU bekommen - automatisch aus Schülerzahlen")
            anschaffung = st.text_input("Anschaffungsdatum (JJJJ-MM-TT)",
                                        value=existing.get("anschaffung",""),
                                        placeholder="2023-08-01")
            bestellbar  = st.checkbox("Im Schulbuchkatalog bestellbar",
                                      value=existing.get("bestellbar", True))
            
            klassensatz = st.checkbox("📚 Als Klassensatz geführt",
                                      value=existing.get("klassensatz", False),
                                      help="Wenn aktiviert: Auf Bestellschein werden 'frei' und 'Kauf' ausgegraut")
            
            # Preis-Feld
            preis_col1, preis_col2 = st.columns([2, 1])
            with preis_col1:
                preis = st.number_input(
                    "Preis (€)", 
                    min_value=0.0, 
                    value=float(existing.get("preis", 0.0)) if existing.get("preis") else 0.0,
                    step=0.50,
                    format="%.2f",
                    help="Manuell eintragen oder per 'Preise aktualisieren' Button holen"
                )
            with preis_col2:
                if existing.get("preis_quelle"):
                    st.caption(f"Quelle: {existing.get('preis_quelle', '')}")
                    if existing.get("preis_aktualisiert"):
                        st.caption(f"Stand: {existing.get('preis_aktualisiert', '')[:10]}")

        st.markdown("---")
        st.markdown("#### 📖 Umlauf-Modus wählen")
        
        # UMLAUF-MODUS (nur wenn NICHT Klassensatz)
        if not klassensatz:
            st.markdown("---")
            st.markdown("### 📚 Umlauf-Modus wählen")
            
            # Radio Buttons für Modus-Auswahl
            modus_idx = {"einzeln": 0, "doppel": 1, "flexibel": 2}.get(existing.get("modus", "einzeln"), 0)
            modus_wahl = st.radio(
                "Wie wird dieses Buch genutzt?",
                options=["⚪ Einzeljahrgang", "🔵 Doppeljahrgang", "🟢 Flexibler Umlauf"],
                index=modus_idx,
                horizontal=True,
                help="Einzeln: Bücher werden jedes Jahr zurückgegeben | Doppel: Bücher werden 2 Jahre behalten | Flexibel: Jahrgangübergreifend mit individueller Rückgabe"
            )
            
            # Modus-String extrahieren
            if "Einzeljahrgang" in modus_wahl:
                modus = "einzeln"
            elif "Doppeljahrgang" in modus_wahl:
                modus = "doppel"
            else:
                modus = "flexibel"
        else:
            # Klassensatz: kein Umlauf nötig
            modus = "klassensatz"
            st.info("📚 **Klassensatz:** Bücher bleiben bei den Lehrkräften - kein Umlauf-Modus nötig!")

        # MODUS-SPEZIFISCHE EINGABEN
        jg1_new = {}
        jg2_new = {}
        uk_new = {}
        flex_new = {}
        
        if modus == "einzeln":
            # ═══ MODUS 1: EINZELJAHRGANG ═══
            st.info("📕 **Einzeljahrgangs-Buch**: Schüler geben das Buch am Ende des Schuljahres zurück.")
            st.caption("Trage ein, wie viele Exemplare aktuell in welcher Klasse im Umlauf sind.")

            klassen_str = st.text_input(
                "Klassen im Umlauf",
                value=", ".join(existing.get("umlauf_klassen", {}).keys()),
                placeholder="z.B. 7a, 7b, 7c",
                key=f"umlauf_klassen_str_{form_id}"
            )
            
            if klassen_str.strip():
                klassen_list = [k.strip() for k in re.split(r"[,;]+", klassen_str) if k.strip()]
                uk_existing = existing.get("umlauf_klassen", {}) or {}
                
                n_cols = min(len(klassen_list), 4)
                cols = st.columns(n_cols)
                
                for i, kl in enumerate(klassen_list):
                    with cols[i % n_cols]:
                        val = st.number_input(
                            f"Klasse {kl}",
                            min_value=0,
                            value=int(uk_existing.get(kl, 0)),
                            key=f"uk_{form_id}_{kl}"
                        )
                        if val > 0:
                            uk_new[kl] = val
            
            # Berechnung anzeigen
            umlauf_summe = sum(uk_new.values())
            verf_next = umlauf_summe + max(lager - RESERVE, 0)
            diff = verf_next - bedarf
            
            st.markdown("---")
            col_calc1, col_calc2, col_calc3 = st.columns(3)
            col_calc1.metric("Im Umlauf", umlauf_summe)
            col_calc2.metric("Verfügbar nächstes Jahr", verf_next,
                           help=f"Alle kommen zurück ({umlauf_summe}) + Lager über Reserve ({max(lager-RESERVE, 0)})")
            col_calc3.metric("Differenz", diff, delta_color="inverse" if diff < 0 else "normal")
            
            if diff < 0:
                st.error(f"⚠️ Es fehlen {abs(diff)} Bücher für nächstes Jahr!")
        
        elif modus == "doppel":
            # ═══ MODUS 2: DOPPELJAHRGANG ═══
            st.info(
                "📘 **Doppeljahrgangs-Buch**: Schüler behalten das Buch 2 Jahre.\n\n"
                "**Beispiel Schuljahr 25/26:**\n"
                "- Jahrgang 1: 5a,5b,5c (bekommen 25/26, behalten bis 26/27)\n"
                "- Jahrgang 2: 6a,6b,6c (haben seit 24/25, geben 26/27 ab)"
            )
            
            col_jg1, col_jg2 = st.columns(2)
            
            # Auto-Generierung aktivieren?
            auto_gen = st.checkbox(
                "🤖 Klassen automatisch generieren",
                value=False,
                help="SEK I: 3-zügig (a,b,c) | SEK II: 2-zügig (.1, .2)"
            )
            
            if auto_gen:
                st.info("💡 Gib nur die Jahrgänge ein (z.B. '5' und '6') - die Klassen werden automatisch erstellt!")
            
            with col_jg1:
                st.markdown("**🟢 Jahrgang 1** (behalten Bücher)")
                jg1_klassen_str = st.text_input(
                    "Klassen Jahrgang 1" if not auto_gen else "Jahrgang 1 (z.B. '5' oder '10g')",
                    value=", ".join(existing.get("jahrgang1_klassen", {}).keys()),
                    placeholder="z.B. 5a, 5b, 5c" if not auto_gen else "z.B. 5 oder 10g",
                    key=f"jg1_klassen_{form_id}"
                )
                
                if jg1_klassen_str.strip():
                    if auto_gen:
                        # Auto-Generierung: Aus "5" mache "5a, 5b, 5c" oder aus "10g" mache "10g.1, 10g.2"
                        jahrgang = jg1_klassen_str.strip()
                        
                        # SEK II: Nur wenn 'g' dabei ist ODER Jahrgang 11/12
                        if jahrgang.endswith('g') or jahrgang in ['11', '12']:
                            # SEK II: 2-zügig mit .1, .2
                            jg1_klassen_list = [f"{jahrgang}.1", f"{jahrgang}.2"]
                        else:
                            # SEK I (5-10): 3-zügig mit a, b, c
                            jg1_klassen_list = [f"{jahrgang}a", f"{jahrgang}b", f"{jahrgang}c"]
                        st.caption(f"Klassen: {', '.join(jg1_klassen_list)}")
                    else:
                        # Manuelle Eingabe
                        jg1_klassen_list = [k.strip() for k in re.split(r"[,;]+", jg1_klassen_str) if k.strip()]
                    
                    jg1_existing = existing.get("jahrgang1_klassen", {}) or {}
                    
                    for kl in jg1_klassen_list:
                        val = st.number_input(
                            f"Klasse {kl}",
                            min_value=0,
                            value=int(jg1_existing.get(kl, 0)),
                            key=f"jg1_{form_id}_{kl}"
                        )
                        # WICHTIG: Auch 0-Werte speichern!
                        jg1_new[kl] = val
            
            with col_jg2:
                st.markdown("**🔴 Jahrgang 2** (geben zurück)")
                jg2_klassen_str = st.text_input(
                    "Klassen Jahrgang 2" if not auto_gen else "Jahrgang 2 (z.B. '6' oder '11')",
                    value=", ".join(existing.get("jahrgang2_klassen", {}).keys()),
                    placeholder="z.B. 6a, 6b, 6c" if not auto_gen else "z.B. 6 oder 11",
                    key=f"jg2_klassen_{form_id}"
                )
                
                if jg2_klassen_str.strip():
                    if auto_gen:
                        # Auto-Generierung
                        jahrgang = jg2_klassen_str.strip()
                        
                        # SEK II: Nur wenn 'g' dabei ist ODER Jahrgang 11/12
                        if jahrgang.endswith('g') or jahrgang in ['11', '12']:
                            # SEK II: 2-zügig
                            jg2_klassen_list = [f"{jahrgang}.1", f"{jahrgang}.2"]
                        else:
                            # SEK I (5-10): 3-zügig
                            jg2_klassen_list = [f"{jahrgang}a", f"{jahrgang}b", f"{jahrgang}c"]
                        st.caption(f"Klassen: {', '.join(jg2_klassen_list)}")
                    else:
                        # Manuelle Eingabe
                        jg2_klassen_list = [k.strip() for k in re.split(r"[,;]+", jg2_klassen_str) if k.strip()]
                    
                    jg2_existing = existing.get("jahrgang2_klassen", {}) or {}
                    
                    for kl in jg2_klassen_list:
                        val = st.number_input(
                            f"Klasse {kl}",
                            min_value=0,
                            value=int(jg2_existing.get(kl, 0)),
                            key=f"jg2_{form_id}_{kl}"
                        )
                        # WICHTIG: Auch 0-Werte speichern!
                        jg2_new[kl] = val
            
            # Berechnung anzeigen
            jg1_summe = sum(jg1_new.values())
            jg2_summe = sum(jg2_new.values())
            verf_next = jg2_summe + max(lager - RESERVE, 0)
            diff = verf_next - bedarf
            
            st.markdown("---")
            col_calc1, col_calc2, col_calc3, col_calc4 = st.columns(4)
            col_calc1.metric("Jahrgang 1 (behalten)", jg1_summe)
            col_calc2.metric("Jahrgang 2 (zurück)", jg2_summe)
            col_calc3.metric("Verfügbar nächstes Jahr", verf_next,
                           help=f"Jahrgang 2 ({jg2_summe}) + Lager über Reserve ({max(lager-RESERVE, 0)})")
            col_calc4.metric("Differenz", diff, delta_color="inverse" if diff < 0 else "normal")
            
            if diff < 0:
                st.error(f"⚠️ Es fehlen {abs(diff)} Bücher für nächstes Jahr!")
        
        else:  # modus == "flexibel"
            # ═══ MODUS 3: FLEXIBLER UMLAUF ═══
            st.info(
                "🟢 **Flexibler Umlauf**: Jahrgangübergreifende Nutzung mit individueller Rückgabe.\n\n"
                "**Beispiel DaZ Band 1:**\n"
                "- 5a: 8 Schüler haben es, 3 geben ab (erreichen Sprachniveau)\n"
                "- 6b: 5 Schüler haben es, 2 geben ab\n"
                "- 7c: 4 Schüler haben es, 4 geben ab"
            )
            
            flex_klassen_str = st.text_input(
                "Klassen im Umlauf",
                value=", ".join(existing.get("flex_klassen", {}).keys()),
                placeholder="z.B. 5a, 6b, 7c, 8a",
                key=f"flex_klassen_{form_id}"
            )
            
            if flex_klassen_str.strip():
                flex_klassen_list = [k.strip() for k in re.split(r"[,;]+", flex_klassen_str) if k.strip()]
                flex_existing = existing.get("flex_klassen", {}) or {}
                
                st.markdown("**Pro Klasse: Wie viele haben das Buch und wie viele geben es zurück?**")
                
                n_cols = min(len(flex_klassen_list), 3)
                cols = st.columns(n_cols)
                
                for i, kl in enumerate(flex_klassen_list):
                    with cols[i % n_cols]:
                        st.markdown(f"**Klasse {kl}**")
                        
                        existing_data = flex_existing.get(kl, {}) if isinstance(flex_existing.get(kl), dict) else {}
                        
                        umlauf = st.number_input(
                            f"Im Umlauf",
                            min_value=0,
                            value=int(existing_data.get("umlauf", 0)),
                            key=f"flex_umlauf_{form_id}_{kl}"
                        )
                        
                        zurueck = st.number_input(
                            f"Davon zurück ↩",
                            min_value=0,
                            value=int(existing_data.get("zurueck", 0)),
                            key=f"flex_zurueck_{form_id}_{kl}",
                            help="Wie viele Schüler geben das Buch nächstes Jahr ab?"
                        )
                        
                        if umlauf > 0:
                            flex_new[kl] = {"umlauf": umlauf, "zurueck": zurueck}
            
            # Berechnung anzeigen
            umlauf_summe = sum(d.get("umlauf", 0) for d in flex_new.values())
            zurueck_summe = sum(d.get("zurueck", 0) for d in flex_new.values())
            behalten_summe = umlauf_summe - zurueck_summe
            verf_next = zurueck_summe + max(lager - RESERVE, 0)
            diff = verf_next - bedarf
            
            st.markdown("---")
            col_calc1, col_calc2, col_calc3, col_calc4, col_calc5 = st.columns(5)
            col_calc1.metric("Im Umlauf gesamt", umlauf_summe)
            col_calc2.metric("Behalten", behalten_summe, help="Schüler die das Buch nächstes Jahr weiter nutzen")
            col_calc3.metric("Kommen zurück ↩", zurueck_summe)
            col_calc4.metric("Verfügbar nächstes Jahr", verf_next,
                           help=f"Zurück ({zurueck_summe}) + Lager über Reserve ({max(lager-RESERVE, 0)})")
            col_calc5.metric("Differenz", diff, delta_color="inverse" if diff < 0 else "normal")
            
            if diff < 0:
                st.error(f"⚠️ Es fehlen {abs(diff)} Bücher für nächstes Jahr!")

        notizen = st.text_area("Notizen", value=existing.get("notizen",""),
                               placeholder='z.B. "DaZ Band 1 für Anfänger"')

        submitted = st.form_submit_button("💾 Speichern", use_container_width=True)

    if submitted:
        if not isbn.strip() or not titel.strip():
            st.error("ISBN und Titel sind Pflichtfelder.")
            return False

        buch = {
            "isbn":           isbn.strip(),
            "titel":          titel.strip(),
            "fach":           fach,
            "verlag":         verlag.strip(),
            "klasse":         klasse.strip(),
            "modus":          modus,
            "lager":          lager,
            "bedarf_next":    bedarf,
            "anschaffung":    anschaffung.strip(),
            "bestellbar":     bestellbar,
            "klassensatz":    klassensatz,
            "notizen":        notizen.strip(),
        }
        
        # Preis nur speichern wenn manuell eingegeben (> 0)
        if preis > 0:
            buch["preis"] = preis
            if not existing.get("preis_quelle"):  # Nur wenn noch keine Quelle da ist
                buch["preis_quelle"] = "Manuell"
        
        # Modus-spezifische Daten
        if modus == "doppel":
            buch["jahrgang1_klassen"] = jg1_new
            buch["jahrgang2_klassen"] = jg2_new
            buch["umlauf_klassen"] = {}
            buch["flex_klassen"] = {}
        elif modus == "flexibel":
            # Validierung: zurück darf nicht größer als umlauf sein
            fehler = []
            for kl, data in flex_new.items():
                if data.get("zurueck", 0) > data.get("umlauf", 0):
                    fehler.append(f"Klasse {kl}: Es können nicht mehr Bücher zurückkommen ({data['zurueck']}) als im Umlauf sind ({data['umlauf']})")
            
            if fehler:
                st.error("❌ Fehler:\n\n" + "\n".join(fehler))
                return False
            
            buch["flex_klassen"] = flex_new
            buch["umlauf_klassen"] = {}
            buch["jahrgang1_klassen"] = {}
            buch["jahrgang2_klassen"] = {}
        else:
            buch["umlauf_klassen"] = uk_new
            buch["jahrgang1_klassen"] = {}
            buch["jahrgang2_klassen"] = {}
            buch["flex_klassen"] = {}
        
        # Legacy-Feld für Kompatibilität
        buch["doppeljahrgang"] = (modus == "doppel")
        
        if save_book(db, buch):
            st.success(f"✅ Buch '{titel}' erfolgreich gespeichert!")
            st.session_state.pop("edit_isbn", None)
            st.session_state["reload"] = True
            return True
    return False


# ═══════════════════════════════════════════════════════════════════════════════
#  HAUPTANWENDUNG
# ═══════════════════════════════════════════════════════════════════════════════

def main_app():
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    db = init_firebase()

    if st.session_state.get("reload") or "buecher" not in st.session_state:
        st.session_state["buecher"] = load_all(db)
        st.session_state.pop("reload", None)

    buecher = st.session_state["buecher"]
    df      = buecher_zu_df(buecher)

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.title(APP_TITLE)
        st.caption(f"Angemeldet als: **{st.session_state.get('user_email','')}**")

        if st.button("🚪 Abmelden", use_container_width=True):
            for k in ["logged_in","user_email","buecher","edit_isbn"]:
                st.session_state.pop(k, None)
            st.rerun()

        st.markdown("---")
        st.subheader("🔍 Filter")
        f_fach   = st.selectbox("Fach", ["Alle"] + sorted(df["fach"].unique().tolist()) if not df.empty else ["Alle"])
        f_klasse = st.selectbox("Klasse", ["Alle"] + ALLE_KLASSEN)
        f_alarm  = st.checkbox("Nur Alarm-Bücher (Nachbestellen)")
        f_text   = st.text_input("Suche (Titel / ISBN)", placeholder="Suche …")

        st.markdown("---")
        st.subheader("📥 Export")

        fname_txt = f"schulbuch_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        # Export mit Schülerzahlen
        txt_data = export_txt(df) if not df.empty else b"Keine Daten"
        if not df.empty:
            schueler_txt = export_schuelerzahlen_txt(db)
            txt_data = txt_data.decode('utf-8') + "\n" + schueler_txt
            txt_data = txt_data.encode('utf-8')
        
        st.download_button(
            "⬇️ Backup als TXT",
            data    = txt_data,
            file_name = fname_txt,
            mime    = "text/plain",
            use_container_width=True,
        )

        if not df.empty:
            # PDF nur bei Klick generieren
            if st.button("📄 PDF generieren", use_container_width=True, key="gen_pdf"):
                with st.spinner("PDF wird erstellt..."):
                    pdf_bytes = export_pdf(df)
                    st.download_button(
                        "⬇️ PDF herunterladen",
                        data      = pdf_bytes,
                        file_name = f"schulbuecher_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime      = "application/pdf",
                        use_container_width=True,
                        key="download_pdf"
                    )

        st.markdown("---")
        st.subheader("💰 Preise")
        
        if not df.empty:
            # Excel-Import
            st.markdown("**📤 Schulbuchkatalog importieren**")
            uploaded_file = st.file_uploader(
                "Katalog.xlsm hochladen",
                type=['xlsm', 'xlsx'],
                help="Lade die Excel-Datei vom Thüringer Schulbuchkatalog hoch",
                key="katalog_upload"
            )
            
            if uploaded_file:
                if st.button("📥 Preise aus Katalog importieren", use_container_width=True):
                    with st.spinner("Importiere Preise aus Katalog..."):
                        result = importiere_preise_aus_excel(db, buecher, uploaded_file)
                        
                        if "error" in result:
                            st.error(f"❌ Fehler beim Import: {result['error']}")
                        elif result["erfolg"] > 0:
                            st.success(
                                f"✅ {result['erfolg']} von {result['gesamt']} Preisen importiert!\n\n"
                                f"Quelle: Schulbuchkatalog {result['katalog_jahr']}"
                            )
                            if result["fehler"] > 0:
                                st.info(f"ℹ️ {result['fehler']} Bücher nicht im Katalog gefunden")
                        else:
                            st.warning("⚠️ Keine passenden ISBNs gefunden.")
                        
                        # Daten neu laden
                        st.session_state["reload"] = True
                        st.rerun()
            
            st.markdown("---")
            
            # Google Books API (als Alternative)
            if st.button("🔄 Preise von Google Books", use_container_width=True, 
                        help="Holt aktuelle Preise von Google Books API (oft nicht verfügbar für Schulbücher)"):
                with st.spinner("Preise werden aktualisiert..."):
                    result = aktualisiere_preise(db, buecher)
                    
                    if result["erfolg"] > 0:
                        st.success(f"✅ {result['erfolg']} von {result['gesamt']} Preisen aktualisiert!")
                    else:
                        st.warning(f"⚠️ Keine Preise gefunden. Nutze besser den Katalog-Import.")
                    
                    if result["fehler"] > 0:
                        st.info(f"ℹ️ {result['fehler']} Bücher: Kein Preis verfügbar")
                    
                    # Daten neu laden
                    st.session_state["reload"] = True
                    st.rerun()
            
            # Preisstatistik
            preise_vorhanden = df[df["preis"].notna() & (df["preis"] > 0)]
            if not preise_vorhanden.empty:
                st.caption(f"📊 {len(preise_vorhanden)} von {len(df)} Büchern haben Preise")
                gesamt_wert = (preise_vorhanden["preis"] * preise_vorhanden["gesamt"]).sum()
                st.caption(f"💶 Gesamtwert: {gesamt_wert:.2f} €")

        st.markdown("---")
        st.subheader("📊 Bestand pro Klasse")
        if not df.empty:
            df_plot = (
                df.groupby("klasse")["gesamt"]
                .sum()
                .reset_index()
                .rename(columns={"klasse":"Klasse","gesamt":"Bücher gesamt"})
            )
            fig = px.bar(
                df_plot, x="Klasse", y="Bücher gesamt",
                color="Bücher gesamt",
                color_continuous_scale="Blues",
                title="Gesamtbestand nach Klasse",
                height=350,
            )
            fig.update_layout(
                showlegend=False,
                margin=dict(l=10,r=10,t=40,b=30),
                plot_bgcolor="rgba(0,0,0,0)",
                xaxis=dict(tickangle=-45),
            )
            st.plotly_chart(fig, use_container_width=True)

        st.markdown("---")
        if st.button("🔄 Daten neu laden", use_container_width=True):
            st.session_state["reload"] = True
            st.rerun()

    # ── Kopfzeile ─────────────────────────────────────────────────────────────
    st.title(APP_TITLE)

    # ── Alarm-Banner ─────────────────────────────────────────────────────────
    alarm_df = df[df["alarm"] == True] if not df.empty else pd.DataFrame()
    if not alarm_df.empty:
        with st.expander(
            f"⚠️ NACHBESTELL-ALARM: {len(alarm_df)} Buch/Bücher müssen nachbestellt werden!",
            expanded=True
        ):
            st.error("Die folgenden Bücher reichen für das nächste Schuljahr **nicht** aus "
                     f"(verfügbar < Bedarf, Mindestlagerbestand {RESERVE} Exemplare eingerechnet):")
            alarm_show = alarm_df[[
                "isbn","titel","fach","klasse","modus",
                "umlauf_gesamt","lager","verfuegbar_next","bedarf_next","differenz",
                "bestellbar","notizen"
            ]].copy()
            alarm_show.columns = [
                "ISBN","Titel","Fach","Klasse","Modus",
                "Umlauf","Lager","Verfügbar nächstes Jahr","Bedarf","Fehlend",
                "Bestellbar","Notizen"
            ]
            st.dataframe(
                alarm_show,
                use_container_width=True,
                hide_index=True,
            )

    # ── Tabs ──────────────────────────────────────────────────────────────────
    tab_liste, tab_neu, tab_edit, tab_details, tab_schueler, tab_etiketten, tab_bestellschein = st.tabs([
        "📋 Bestandsliste",
        "➕ Buch hinzufügen",
        "✏️ Buch bearbeiten / löschen",
        "🔎 Buch-Details",
        "👥 Schülerzahlen",
        "🏷️ Etiketten",
        "📝 Bestellscheine",
    ])

    with tab_liste:
        st.subheader("Aktueller Buchbestand")

        df_view = df.copy() if not df.empty else df
        if not df_view.empty:
            if f_fach != "Alle":
                df_view = df_view[df_view["fach"] == f_fach]
            if f_klasse != "Alle":
                df_view = df_view[
                    df_view["klasse"].str.contains(f_klasse, na=False) |
                    df_view["umlauf_klassen"].str.contains(f_klasse, na=False)
                ]
            if f_alarm:
                df_view = df_view[df_view["alarm"] == True]
            if f_text:
                mask = (
                    df_view["titel"].str.contains(f_text, case=False, na=False) |
                    df_view["isbn"].str.contains(f_text, case=False, na=False)
                )
                df_view = df_view[mask]

        if df_view.empty:
            st.info("Keine Bücher gefunden. Füge über den Tab '➕ Buch hinzufügen' dein erstes Buch hinzu.")
        else:
            show_cols = [
                "isbn","titel","fach","klasse","modus",
                "umlauf_klassen","umlauf_gesamt","lager","gesamt",
                "bedarf_next","verfuegbar_next","differenz",
                "preis","anschaffung","bestellbar","notizen"
            ]
            display_df = df_view[show_cols].copy()
            
            # Preis formatieren + manuelle Preise markieren
            def format_preis(row):
                preis = row["preis"]
                preis_quelle = row.get("preis_quelle", "")
                
                if pd.notna(preis) and preis > 0:
                    preis_str = f"{preis:.2f} €".replace(".", ",")
                    # Markiere manuelle Preise
                    if preis_quelle == "Manuell":
                        return f"*{preis_str}"  # Stern für manuell
                    return preis_str
                return "-"
            
            display_df["preis_formatted"] = df_view.apply(format_preis, axis=1)
            display_df = display_df.drop("preis", axis=1)
            
            display_df.columns = [
                "ISBN","Titel","Fach","Klasse","Modus",
                "Umlauf pro Klasse","Umlauf Σ","Lager","Gesamt",
                "Bedarf n.J.","Verfügbar n.J.","Differenz",
                "Anschaffung","Bestellbar","Notizen","Preis"
            ]
            
            # Spalten neu ordnen - Preis nach Differenz
            cols = ["ISBN","Titel","Fach","Klasse","Modus",
                   "Umlauf pro Klasse","Umlauf Σ","Lager","Gesamt",
                   "Bedarf n.J.","Verfügbar n.J.","Differenz","Preis",
                   "Anschaffung","Bestellbar","Notizen"]
            display_df = display_df[cols]

            def style_row(row):
                styles = [""] * len(row)
                if row["Differenz"] < 0:
                    styles = ["background-color: #ffe0e0"] * len(row)
                elif row["Differenz"] < 5:
                    styles = ["background-color: #fff3cd"] * len(row)
                
                # Preis kursiv wenn manuell (markiert mit *)
                preis_idx = list(row.index).index("Preis")
                if str(row["Preis"]).startswith("*"):
                    styles[preis_idx] = styles[preis_idx] + "; font-style: italic"
                
                return styles

            styled = display_df.style.apply(style_row, axis=1)
            
            st.info("💡 **Tipp:** Manuelle Preise sind mit * markiert und kursiv. " 
                   "Zum Bearbeiten eines Buches: Wähle es im Tab '✏️ Buch bearbeiten'")
            
            st.dataframe(styled, use_container_width=True, hide_index=True, height=500)

            col_a, col_b, col_c = st.columns(3)
            col_a.metric("Bücher gesamt", len(df_view))
            col_b.metric("Exemplare gesamt", int(df_view["gesamt"].sum()))
            col_c.metric("🔴 Nachbestellen",
                         int(df_view["alarm"].sum()),
                         delta_color="inverse")
            
            # Schnellauswahl zum Bearbeiten
            st.markdown("---")
            st.markdown("**⚡ Schnellzugriff:**")
            
            buch_optionen_quick = {
                f"{b['titel']} - {b['fach']} (ISBN: {b['isbn']})": b['isbn']
                for b in buecher
            }
            
            col_select, col_hint = st.columns([3, 2])
            
            with col_select:
                selected_quick = st.selectbox(
                    "Buch vorauswählen:",
                    [""] + list(buch_optionen_quick.keys()),
                    key="quick_select",
                    label_visibility="collapsed"
                )
            
            with col_hint:
                if selected_quick:
                    st.success("✅ Vorausgewählt!")
                    st.caption("→ Wechsle zum Tab  \n**✏️ Buch bearbeiten**")
            
            if selected_quick:
                st.session_state["edit_isbn"] = buch_optionen_quick[selected_quick]

    with tab_neu:
        buch_formular(db, existing=None)

    with tab_edit:
        st.subheader("Buch bearbeiten oder löschen")

        if df.empty:
            st.info("Noch keine Bücher vorhanden.")
        else:
            buch_optionen = {
                f"{b['titel']} (ISBN: {b['isbn']})": b
                for b in buecher
            }
            
            # Vorauswahl aus Schnellzugriff?
            default_index = 0
            if "edit_isbn" in st.session_state and st.session_state["edit_isbn"]:
                # Finde Index des vorausgewählten Buches
                for idx, (label, buch) in enumerate(buch_optionen.items()):
                    if buch["isbn"] == st.session_state["edit_isbn"]:
                        default_index = idx
                        break
            
            auswahl_label = st.selectbox(
                "Buch auswählen",
                list(buch_optionen.keys()),
                index=default_index,
                key="edit_select"
            )
            gewaehltes_buch = buch_optionen[auswahl_label]

            col_edit, col_del = st.columns([3, 1])
            with col_edit:
                st.markdown("#### ✏️ Daten bearbeiten")
                buch_formular(db, existing=gewaehltes_buch)

            with col_del:
                st.markdown("#### 🗑️ Buch löschen")
                st.warning(
                    f"**{gewaehltes_buch['titel']}**\n\n"
                    f"ISBN: {gewaehltes_buch['isbn']}\n\n"
                    "Diese Aktion kann nicht rückgängig gemacht werden!"
                )
                confirm = st.checkbox(
                    "Ja, ich möchte dieses Buch unwiderruflich löschen",
                    key="confirm_delete"
                )
                if st.button("🗑️ Endgültig löschen", disabled=not confirm,
                             use_container_width=True, type="primary"):
                    if delete_book(db, gewaehltes_buch["isbn"]):
                        st.success(f"✅ '{gewaehltes_buch['titel']}' wurde gelöscht.")
                        st.session_state["reload"] = True
                        st.rerun()

    with tab_details:
        st.subheader("🔎 Detailansicht")

        if df.empty:
            st.info("Noch keine Bücher vorhanden.")
        else:
            detail_optionen = {
                f"{b['titel']} (ISBN: {b['isbn']})": b
                for b in buecher
            }
            detail_label = st.selectbox(
                "Buch auswählen",
                list(detail_optionen.keys()),
                key="detail_select"
            )
            b = berechne_felder(detail_optionen[detail_label].copy())

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Umlauf gesamt", b["umlauf_gesamt"])
            col2.metric("Lager", b["lager"], help=f"Mindestlagerbestand: {RESERVE} Exemplare")
            col3.metric("Gesamtbestand", b["gesamt"])
            col4.metric(
                "Verfügbar nächstes Jahr",
                b["verfuegbar_next"],
                delta=b["differenz"],
                delta_color="normal",
                help="Bücher die zurückkommen + Lager über Reserve"
            )

            if b["alarm"]:
                st.error(
                    f"⚠️ **NACHBESTELLBEDARF**: Es fehlen **{abs(b['differenz'])} Exemplare**! "
                    f"(Bedarf: {b['bedarf_next']}, Verfügbar: {b['verfuegbar_next']})\n\n"
                    f"Bestellbar im Katalog: **{'Ja' if b['bestellbar'] else 'Nein'}**"
                )
            elif b["differenz"] < RESERVE:
                st.warning(
                    f"🟡 Knappe Reserve: nur {b['differenz']} Exemplare über Bedarf."
                )
            else:
                st.success(
                    f"✅ Bestand ausreichend. Überschuss: {b['differenz']} Exemplare."
                )

            st.markdown("#### Umlauf-Details")
            
            modus = b.get("modus", "einzeln")
            
            if modus == "flexibel":
                st.info("🟢 **Flexibler Umlauf** (jahrgangübergreifend, individuelle Rückgabe)")
                flex = b.get("flex_klassen", {}) or {}
                if flex:
                    flex_data = []
                    for kl, data in sorted(flex.items()):
                        if isinstance(data, dict):
                            flex_data.append({
                                "Klasse": kl,
                                "Im Umlauf": data.get("umlauf", 0),
                                "Kommen zurück": data.get("zurueck", 0),
                                "Behalten": data.get("umlauf", 0) - data.get("zurueck", 0)
                            })
                    if flex_data:
                        flex_df = pd.DataFrame(flex_data)
                        st.dataframe(flex_df, hide_index=True, use_container_width=True)
                        
                        col_sum1, col_sum2, col_sum3 = st.columns(3)
                        col_sum1.metric("Gesamt im Umlauf", sum(d["Im Umlauf"] for d in flex_data))
                        col_sum2.metric("Kommen zurück ↩", b.get("zurueck_gesamt", 0))
                        col_sum3.metric("Behalten", sum(d["Behalten"] for d in flex_data))
                else:
                    st.caption("Keine Klassen eingetragen")
            
            elif modus == "doppel":
                st.info("📘 **Doppeljahrgangs-Buch** (wird 2 Jahre behalten)")
                
                col_jg1, col_jg2 = st.columns(2)
                
                with col_jg1:
                    st.markdown("**🟢 Jahrgang 1** (behalten Bücher)")
                    jg1 = b.get("jahrgang1_klassen") or {}
                    if jg1:
                        jg1_df = pd.DataFrame(
                            [(k, v) for k, v in sorted(jg1.items())],
                            columns=["Klasse", "Exemplare"]
                        )
                        st.dataframe(jg1_df, hide_index=True, use_container_width=True)
                        st.metric("Summe Jahrgang 1", b.get("jahrgang1_gesamt", 0))
                    else:
                        st.caption("Keine Klassen eingetragen")
                
                with col_jg2:
                    st.markdown("**🔴 Jahrgang 2** (geben zurück)")
                    jg2 = b.get("jahrgang2_klassen") or {}
                    if jg2:
                        jg2_df = pd.DataFrame(
                            [(k, v) for k, v in sorted(jg2.items())],
                            columns=["Klasse", "Exemplare"]
                        )
                        st.dataframe(jg2_df, hide_index=True, use_container_width=True)
                        st.metric("Summe Jahrgang 2", b.get("jahrgang2_gesamt", 0))
                    else:
                        st.caption("Keine Klassen eingetragen")
            
            else:
                st.info("📕 **Einzeljahrgangs-Buch**")
                uk = b.get("umlauf_klassen") or {}
                if uk:
                    uk_df = pd.DataFrame(
                        [(k, v) for k, v in sorted(uk.items())],
                        columns=["Klasse", "Exemplare"]
                    )
                    st.dataframe(uk_df, hide_index=True, use_container_width=True)

            st.markdown("#### 📋 Buchdetails")
            meta_col1, meta_col2 = st.columns(2)
            with meta_col1:
                st.markdown(f"**Fach:** {b.get('fach','')}")
                st.markdown(f"**Klasse(n):** {b.get('klasse','')}")
                modus_namen = {'einzeln':'Einzeljahrgang','doppel':'Doppeljahrgang','flexibel':'Flexibler Umlauf'}
                st.markdown(f"**Modus:** {modus_namen.get(modus, modus)}")
                st.markdown(f"**ISBN:** {b.get('isbn','')}")
            with meta_col2:
                st.markdown(f"**Anschaffung:** {b.get('anschaffung','')}")
                
                # Preis anzeigen
                preis = b.get('preis')
                if preis and preis > 0:
                    preis_quelle = b.get('preis_quelle', '')
                    preis_datum = b.get('preis_aktualisiert', '')
                    preis_info = f"{preis:.2f} €"
                    if preis_quelle:
                        preis_info += f" ({preis_quelle}"
                        if preis_datum:
                            preis_info += f", {preis_datum[:10]}"
                        preis_info += ")"
                    st.markdown(f"**Preis:** {preis_info}")
                else:
                    st.markdown("**Preis:** Nicht verfügbar")
                
                st.markdown(f"**Im Katalog bestellbar:** {'✅ Ja' if b.get('bestellbar') else '❌ Nein'}")
                st.markdown(f"**Notizen:** {b.get('notizen','–')}")
    
    # ══════════════════════════════════════════════════════════════════════════
    
    # ══════════════════════════════════════════════════════════════════════════
    # TAB: SCHÜLERZAHLEN
    # ══════════════════════════════════════════════════════════════════════════
    
    with tab_schueler:
        st.subheader("👥 Schülerzahlen verwalten")
        
        st.info("💡 **Schülerzahlen für automatischen Bedarf**\n\n"
                "Gib hier die Schülerzahlen pro Jahrgang ein. Diese werden automatisch als Bedarf "
                "bei neu angelegten Büchern übernommen.")
        
        # Schuljahr wählen
        from datetime import datetime
        aktuelles_jahr = datetime.now().year
        aktueller_monat = datetime.now().month
        
        # Schuljahr berechnen (ab August = neues Schuljahr)
        if aktueller_monat >= 8:
            schuljahr_start = aktuelles_jahr
        else:
            schuljahr_start = aktuelles_jahr - 1
        schuljahr_end = schuljahr_start + 1
        schuljahr_default = f"{schuljahr_start}/{schuljahr_end}"
        
        col_sj, col_btn = st.columns([3, 1])
        with col_sj:
            schuljahr_input = st.text_input("📅 Schuljahr", value=schuljahr_default)
        
        # Lade bestehende Schülerzahlen
        try:
            schueler_ref = db.collection('schuelerzahlen').document(schuljahr_input)
            schueler_doc = schueler_ref.get()
            if schueler_doc.exists:
                schueler_data = schueler_doc.to_dict()
            else:
                schueler_data = {}
        except Exception as e:
            st.error(f"Fehler beim Laden: {e}")
            schueler_data = {}
        
        st.markdown("---")
        st.markdown("### Schülerzahlen pro Jahrgang")
        
        jahrgaenge = ["5", "6", "7", "8", "9", "10", "10g", "11", "12"]
        neue_daten = {}
        
        for jg in jahrgaenge:
            st.markdown(f"**Klasse {jg}**")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                gesamt = st.number_input(
                    f"Gesamt",
                    min_value=0,
                    value=schueler_data.get(f"{jg}_gesamt", 0),
                    key=f"schueler_{jg}_gesamt",
                    help="Anzahl Schüler gesamt (für Standard-Fächer)"
                )
                neue_daten[f"{jg}_gesamt"] = gesamt
            
            with col2:
                spanisch = st.number_input(
                    f"Spanisch",
                    min_value=0,
                    value=schueler_data.get(f"{jg}_spanisch", 0),
                    key=f"schueler_{jg}_spanisch"
                )
                neue_daten[f"{jg}_spanisch"] = spanisch
            
            with col3:
                religion = st.number_input(
                    f"Religion",
                    min_value=0,
                    value=schueler_data.get(f"{jg}_religion", 0),
                    key=f"schueler_{jg}_religion"
                )
                neue_daten[f"{jg}_religion"] = religion
            
            with col4:
                daz = st.number_input(
                    f"DaZ",
                    min_value=0,
                    value=schueler_data.get(f"{jg}_daz", 0),
                    key=f"schueler_{jg}_daz"
                )
                neue_daten[f"{jg}_daz"] = daz
        
        st.markdown("---")
        
        if st.button("💾 Schülerzahlen speichern", type="primary"):
            try:
                schueler_ref.set(neue_daten)
                st.success(f"✅ Schülerzahlen für {schuljahr_input} gespeichert!")
            except Exception as e:
                st.error(f"❌ Fehler beim Speichern: {e}")
    
    # ══════════════════════════════════════════════════════════════════════════
    # TAB: ETIKETTEN
    # ══════════════════════════════════════════════════════════════════════════
    
    with tab_etiketten:
        st.subheader("🏷️ Buchetiketten erstellen")
        
        st.info("💡 **Erstelle Etiketten für neue Bücher!**\n\n"
                "AVERY Zweckform 3424 - 12 Etiketten pro A4 Seite (105 x 48 mm)")
        
        if df.empty:
            st.warning("Noch keine Bücher vorhanden.")
        else:
            # Buch auswählen
            buecher_sortiert = sortiere_buecher_fach_jahrgang(buecher)
            
            buch_optionen = {
                f"{b['fach']} {b.get('klasse', '')} - {b['titel']}": b
                for b in buecher_sortiert
            }
            
            selected_buch_label = st.selectbox(
                "Buch auswählen:",
                list(buch_optionen.keys())
            )
            
            selected_buch = buch_optionen[selected_buch_label]
            
            col1, col2 = st.columns(2)
            
            with col1:
                titel_etikett = st.text_input("Titel", value=selected_buch.get("titel", ""))
                klasse_etikett = st.text_input("Klasse", value=selected_buch.get("klasse", ""))
                fach_etikett = st.text_input("Fach", value=selected_buch.get("fach", ""))
                
                # Anschaffungsdatum
                heute = datetime.now()
                anschaffung_datum = st.date_input("Anschaffungsdatum", value=heute)
                anschaffung_str = anschaffung_datum.strftime("%d.%m.%Y")
            
            with col2:
                # Neuanschaffung oder Ergänzung
                typ = st.radio(
                    "Art der Anschaffung:",
                    ["Neuanschaffung", "Ergänzung"],
                    horizontal=True
                )
                
                if typ == "Neuanschaffung":
                    start_nr = 1
                else:
                    # Berechne aus Lager + Umlauf
                    lager = selected_buch.get("lager", 0)
                    
                    # Berechne Umlauf je nach Modus
                    modus = selected_buch.get("modus", "einzeln")
                    umlauf_gesamt = 0
                    
                    if modus == "einzeln":
                        klassen_dict = selected_buch.get("klassen_umlauf", {})
                        umlauf_gesamt = sum(klassen_dict.values())
                    elif modus == "doppel":
                        jg1 = selected_buch.get("jahrgang1_klassen", {})
                        jg2 = selected_buch.get("jahrgang2_klassen", {})
                        umlauf_gesamt = sum(jg1.values()) + sum(jg2.values())
                    elif modus == "flexibel":
                        flex_dict = selected_buch.get("flexibel_umlauf", {})
                        umlauf_gesamt = sum(flex_dict.values())
                    
                    start_nr = lager + umlauf_gesamt + 1
                
                start_nummer = st.number_input(
                    "Startnummer (#)",
                    min_value=1,
                    value=start_nr,
                    help="Erste Nummer für die Etiketten"
                )
                
                anzahl = st.number_input(
                    "Anzahl Etiketten",
                    min_value=1,
                    max_value=500,
                    value=10,
                    help="Wie viele Etiketten sollen erstellt werden?"
                )
            
            st.markdown("---")
            
            # Logo laden
            logo_bytes_etikett = None
            logo_pfade = [
                "logo.jpg", "./logo.jpg", 
                "logo.png", "./logo.png",
                "Logo.jpg", "./Logo.jpg",
                "Logo.png", "./Logo.png",
            ]
            for pfad in logo_pfade:
                try:
                    with open(pfad, "rb") as f:
                        logo_bytes_etikett = f.read()
                        break
                except:
                    continue
            
            col_docx, col_pdf = st.columns(2)
            
            with col_docx:
                if st.button("📥 Als DOCX herunterladen", use_container_width=True):
                    try:
                        docx_bytes = generiere_etiketten_docx(
                            titel=titel_etikett,
                            klasse=klasse_etikett,
                            fach=fach_etikett,
                            anschaffung=anschaffung_str,
                            start_nr=start_nummer,
                            anzahl=anzahl,
                            logo_bytes=logo_bytes_etikett
                        )
                        
                        filename = f"Etiketten_{fach_etikett}_{klasse_etikett}_{anschaffung_datum.strftime('%Y%m%d')}.docx"
                        
                        st.download_button(
                            label="⬇️ DOCX herunterladen",
                            data=docx_bytes,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Fehler beim Erstellen: {e}")
            
            with col_pdf:
                st.info("💡 PDF-Export: Drucke das DOCX als PDF oder nutze 'Drucken → Als PDF speichern'")
    
    # ══════════════════════════════════════════════════════════════════════════
    # TAB: BESTELLSCHEINE
    # ══════════════════════════════════════════════════════════════════════════
    
    with tab_bestellschein:
        st.subheader("📝 Schulbuchbestellschein erstellen")
        
        st.info("💡 **Erstelle Bestellscheine für deine Schüler!**\n\n"
                "Wähle Bücher aus dem Bestand und generiere ein Word-Dokument.")
        
        # Logo Upload
        logo_bytes = None
        
        # Versuche Logo von GitHub/lokal zu laden - verschiedene Namen
        logo_pfade = [
            "logo.jpg", "./logo.jpg", 
            "logo.png", "./logo.png",
            "Logo.jpg", "./Logo.jpg",
            "Logo.png", "./Logo.png",
            "ginkgo.jpg", "./ginkgo.jpg",
            "Ginkgo.jpg", "./Ginkgo.jpg",
            "schullogo.jpg", "./schullogo.jpg",
            "schullogo.png", "./schullogo.png"
        ]
        for pfad in logo_pfade:
            try:
                with open(pfad, "rb") as f:
                    logo_bytes = f.read()
                    st.success(f"✅ Logo automatisch geladen: {pfad}")
                    break
            except:
                continue
        
        # Falls nicht gefunden: Upload-Option
        if not logo_bytes:
            st.info("💡 **Tipp:** Benenne dein Logo in GitHub als `logo.jpg` oder `logo.png`, dann wird es automatisch geladen!")
            logo_file = st.file_uploader("📷 Logo hochladen", 
                                          type=["jpg", "jpeg", "png"],
                                          help="Lade dein Schullogo hoch - wird rechts oben im Bestellschein angezeigt")
            logo_bytes = logo_file.read() if logo_file else None
        
        col1, col2 = st.columns(2)
        
        # Aktuelles Schuljahr berechnen
        heute = datetime.now()
        jahr = heute.year
        # Schuljahr beginnt im August
        if heute.month >= 8:
            schuljahr_default = f"{jahr}/{jahr+1}"
        else:
            schuljahr_default = f"{jahr-1}/{jahr}"
        
        with col1:
            klassenstufe = st.selectbox("Klassenstufe", 
                                        ["5", "6", "7", "8", "9", "10", "10g", "11", "12"])
            schuljahr = st.text_input("Schuljahr", value=schuljahr_default)
        
        with col2:
            arbeitsheft_leerzeilen = st.number_input("Leerzeilen Arbeitshefte", 
                                                      min_value=0, max_value=20, value=5)
            weiterfuehrung_leerzeilen = st.number_input("Leerzeilen Weiterführung", 
                                                          min_value=0, max_value=10, 
                                                          value=0 if klassenstufe == "5" else 3)
        
        st.markdown("---")
        st.markdown("### 📚 Lehrbücher (aus Bestand)")
        
        # Bücher aus Bestand für diese Klasse
        verfuegbare_buecher = [b for b in buecher if klassenstufe in str(b.get("klasse", ""))]
        
        if not verfuegbare_buecher:
            st.warning(f"Keine Bücher für Klasse {klassenstufe} im Bestand gefunden!")
            ausgewaehlte_buecher = []
        else:
            st.write(f"✅ {len(verfuegbare_buecher)} Bücher für Klasse {klassenstufe} verfügbar")
            
            # Sortiere nach Fach, dann Jahrgang
            verfuegbare_buecher_sortiert = sortiere_buecher_fach_jahrgang(verfuegbare_buecher)
            
            ausgewaehlte_isbns = st.multiselect(
                "Wähle Lehrbücher aus:",
                options=[b["isbn"] for b in verfuegbare_buecher_sortiert],
                format_func=lambda isbn: next((f"{b['fach']} {b.get('klasse', '')} - {b['titel']}" for b in verfuegbare_buecher_sortiert if b['isbn'] == isbn), isbn)
            )
            
            ausgewaehlte_buecher = [b for b in verfuegbare_buecher if b["isbn"] in ausgewaehlte_isbns]
        
        st.markdown("---")
        st.markdown("### 📖 Arbeitshefte (ISBNs eingeben)")
        
        arbeitshefte_text = st.text_area(
            "ISBNs der Arbeitshefte (eine pro Zeile):",
            height=150,
            help="Gib die ISBNs der Arbeitshefte ein - Titel und Verlag werden automatisch aus dem Katalog geholt"
        )
        
        st.markdown("---")
        st.markdown("### 🔄 Weiterführung aus Vorjahr")
        
        if klassenstufe == "5":
            st.info("ℹ️ Klasse 5 hat keine Weiterführung (kommen aus Grundschule)")
            weiterfuehrung_buecher = []
            vorjahr_klasse = None
        else:
            vorjahr_klasse = str(int(klassenstufe.replace('g', '')) - 1)
            if klassenstufe.endswith('g'):
                vorjahr_klasse += 'g'
            
            # Bücher aus dem Vorjahr-Bestand
            weiterfuehrung_verfuegbar = [b for b in buecher if vorjahr_klasse in str(b.get("klasse", ""))]
            
            if not weiterfuehrung_verfuegbar:
                st.info(f"ℹ️ Keine Bücher für Klasse {vorjahr_klasse} im Bestand gefunden (Vorjahr)")
                weiterfuehrung_buecher = []
            else:
                st.write(f"✅ {len(weiterfuehrung_verfuegbar)} Bücher aus Klasse {vorjahr_klasse} verfügbar")
                
                # Sortiere nach Fach, dann Jahrgang
                weiterfuehrung_sortiert = sortiere_buecher_fach_jahrgang(weiterfuehrung_verfuegbar)
                
                weiterfuehrung_isbns = st.multiselect(
                    f"Wähle weitergeführte Bücher aus Klasse {vorjahr_klasse}:",
                    options=[b["isbn"] for b in weiterfuehrung_sortiert],
                    format_func=lambda isbn: next((f"{b['fach']} {b.get('klasse', '')} - {b['titel']}" for b in weiterfuehrung_sortiert if b['isbn'] == isbn), isbn),
                    help="Bücher die vom Vorjahr weitergenutzt werden"
                )
                
                weiterfuehrung_buecher = [b for b in weiterfuehrung_verfuegbar if b["isbn"] in weiterfuehrung_isbns]
        
        st.markdown("---")
        
        # GENERIEREN Button
        if st.button("📄 Bestellschein generieren", type="primary"):
            if not ausgewaehlte_buecher and not arbeitshefte_text.strip():
                st.error("⚠️ Bitte wähle mindestens ein Buch aus oder gib ISBNs ein!")
            else:
                with st.spinner("Erstelle Bestellschein..."):
                    # Prepare Lehrbücher
                    lehrbuecher_data = []
                    for b in ausgewaehlte_buecher:
                        # Hole fehlende Daten aus Katalog
                        katalog_data = hole_buch_aus_katalog(b.get("isbn", ""))
                        
                        # Verlag: Erst aus DB, dann aus Katalog
                        verlag = b.get("verlag", "")
                        if not verlag and katalog_data:
                            verlag = katalog_data.get('verlag', '')
                        
                        # Titel: Aus DB (sollte vorhanden sein)
                        titel = b.get("titel", "")
                        if not titel and katalog_data:
                            titel = katalog_data.get('titel', '')
                        
                        # Preis formatieren
                        preis_str = ""
                        if b.get('preis'):
                            preis_str = f"{b.get('preis', 0):.2f} €".replace(".", ",")
                        elif katalog_data and katalog_data.get('preis'):
                            preis_str = katalog_data.get('preis', '')
                        
                        lehrbuecher_data.append({
                            "fach": b.get("fach", ""),
                            "titel": titel,
                            "verlag": verlag,
                            "isbn": b.get("isbn", ""),
                            "preis": preis_str,
                            "klassensatz": b.get("klassensatz", False)  # Aus DB holen!
                        })
                    
                    # Prepare Arbeitshefte - hole Daten aus Katalog + Google Books
                    arbeitshefte_data = []
                    if arbeitshefte_text.strip():
                        ah_isbns = [isbn.strip() for isbn in arbeitshefte_text.strip().split("\n") if isbn.strip()]
                        for isbn in ah_isbns:
                            # Erst aus Bestand versuchen
                            buch = next((b for b in buecher if b["isbn"] == isbn), None)
                            
                            # Dann aus Katalog
                            katalog_data = hole_buch_aus_katalog(isbn)
                            
                            # Wenn weder Bestand noch Katalog: Google Books
                            google_data = None
                            if not buch and not katalog_data:
                                google_data = hole_preis_von_google_books(isbn)
                            
                            if buch:
                                # Aus Bestand, ergänze mit Katalog
                                verlag = buch.get("verlag", "")
                                if not verlag and katalog_data:
                                    verlag = katalog_data.get('verlag', '')
                                
                                preis_str = f"{buch.get('preis', 0):.2f} €".replace(".", ",") if buch.get('preis') else ""
                                if not preis_str and katalog_data:
                                    preis_str = katalog_data.get('preis', '')
                                
                                arbeitshefte_data.append({
                                    "fach": buch.get("fach", ""),
                                    "titel": buch.get("titel", ""),
                                    "verlag": verlag,
                                    "isbn": isbn,
                                    "preis": preis_str
                                })
                            elif katalog_data:
                                # Nur aus Katalog
                                arbeitshefte_data.append({
                                    "fach": katalog_data.get("fach", ""),
                                    "titel": katalog_data.get("titel", ""),
                                    "verlag": katalog_data.get("verlag", ""),
                                    "isbn": isbn,
                                    "preis": katalog_data.get("preis", "")
                                })
                            elif google_data:
                                # Aus Google Books - versuche Fach zu raten
                                titel = google_data.get("titel", "")
                                fach_geraten = ""
                                # Versuche Fach aus Titel zu erraten
                                titel_lower = titel.lower()
                                if "mathemat" in titel_lower or "mathe" in titel_lower:
                                    fach_geraten = "Mathematik"
                                elif "deutsch" in titel_lower:
                                    fach_geraten = "Deutsch"
                                elif "engl" in titel_lower:
                                    fach_geraten = "Englisch"
                                elif "geo" in titel_lower:
                                    fach_geraten = "Geographie"
                                elif "bio" in titel_lower:
                                    fach_geraten = "Biologie"
                                elif "chem" in titel_lower:
                                    fach_geraten = "Chemie"
                                elif "phys" in titel_lower:
                                    fach_geraten = "Physik"
                                elif "gesch" in titel_lower:
                                    fach_geraten = "Geschichte"
                                # Wenn nicht gefunden: leer lassen
                                
                                preis_str = f"{google_data.get('preis', 0):.2f} €".replace(".", ",") if google_data.get('preis') else ""
                                
                                arbeitshefte_data.append({
                                    "fach": fach_geraten,
                                    "titel": titel,
                                    "verlag": google_data.get("verlag", ""),
                                    "isbn": isbn,
                                    "preis": preis_str
                                })
                            else:
                                # Nichts gefunden - leer lassen
                                arbeitshefte_data.append({
                                    "fach": "",
                                    "titel": "",
                                    "verlag": "",
                                    "isbn": isbn,
                                    "preis": ""
                                })
                    
                    # Prepare Weiterführung
                    weiterfuehrung_data = []
                    if klassenstufe != "5" and weiterfuehrung_buecher:
                        for b in weiterfuehrung_buecher:
                            # Hole fehlende Daten aus Katalog
                            katalog_data = hole_buch_aus_katalog(b.get("isbn", ""))
                            
                            # Verlag: Erst aus DB, dann aus Katalog
                            verlag = b.get("verlag", "")
                            if not verlag and katalog_data:
                                verlag = katalog_data.get('verlag', '')
                            
                            # Titel: Aus DB
                            titel = b.get("titel", "")
                            if not titel and katalog_data:
                                titel = katalog_data.get('titel', '')
                            
                            weiterfuehrung_data.append({
                                "fach": b.get("fach", ""),
                                "titel": titel,
                                "verlag": verlag,
                                "isbn": b.get("isbn", ""),
                                "aus_klasse": vorjahr_klasse
                            })
                    
                    try:
                        # Generiere DOCX
                        docx_bytes = generiere_bestellschein_bytes(
                            klassenstufe=klassenstufe,
                            schuljahr=schuljahr,
                            lehrbuecher=lehrbuecher_data,
                            arbeitshefte=arbeitshefte_data,
                            arbeitsheft_leerzeilen=arbeitsheft_leerzeilen,
                            weiterfuehrung=weiterfuehrung_data,
                            weiterfuehrung_leerzeilen=weiterfuehrung_leerzeilen,
                            logo_bytes=logo_bytes
                        )
                        
                        st.success("✅ Bestellschein erfolgreich erstellt!")
                        
                        # Download Button
                        st.download_button(
                            label="📥 Bestellschein herunterladen",
                            data=docx_bytes,
                            file_name=f"Bestellschein_Klasse_{klassenstufe}_{schuljahr.replace('/', '-')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                    except Exception as e:
                        st.error(f"❌ Fehler beim Erstellen: {str(e)}")
                        st.exception(e)


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def run():
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False

    if not st.session_state["logged_in"]:
        render_login_page()
    else:
        main_app()


if __name__ == "__main__":
    run()
